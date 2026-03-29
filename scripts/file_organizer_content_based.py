#!/usr/bin/env python3
"""
Content-Based Intelligent File Organizer using Schema.org metadata and OCR.

Organizes files based on their actual content rather than just file type.
Uses OCR to extract text from images and PDFs, then classifies by content.
"""

import sys
import os
import shutil
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Any
from collections import defaultdict
from urllib.parse import quote
from contextlib import nullcontext

# OCR and PDF imports
try:
    import pytesseract
    from PIL import Image
    import pypdf
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True

    # HEIC support for OCR
    try:
        from pillow_heif import register_heif_opener
        register_heif_opener()
    except ImportError:
        pass  # HEIC support optional for OCR
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR libraries not available. Install pytesseract, Pillow, pypdf, pdf2image")

# Word document imports
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install python-docx")

# Excel imports
try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("Warning: openpyxl not available. Install openpyxl")

# Add src directory to path (portable)
sys.path.insert(0, str(Path(__file__).parent.parent / 'src'))

from generators import (
    DocumentGenerator,
    ImageGenerator,
    VideoGenerator,
    AudioGenerator,
    CodeGenerator,
    DatasetGenerator,
    ArchiveGenerator
)
from base import PropertyType
from enrichment import MetadataEnricher
from validator import SchemaValidator
from integration import SchemaRegistry

# Graph storage imports
try:
    from storage.graph_store import GraphStore
    from storage.models import File as FileModel, FileStatus
    GRAPH_STORE_AVAILABLE = True
except ImportError:
    GRAPH_STORE_AVAILABLE = False
    print("Warning: GraphStore not available. Database persistence disabled.")

# Image content analysis imports
try:
    from transformers import CLIPProcessor, CLIPModel
    import torch
    import cv2
    VISION_AVAILABLE = True
except ImportError:
    VISION_AVAILABLE = False
    print("Warning: Vision libraries not available. Install transformers, torch, opencv-python")

# Image metadata imports
try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    import piexif
    from geopy.geocoders import Nominatim
    from geopy.exc import GeocoderTimedOut, GeocoderServiceError
    METADATA_AVAILABLE = True

    # HEIC support
    try:
        from pillow_heif import register_heif_opener
        register_heif_opener()
        HEIC_AVAILABLE = True
    except ImportError:
        HEIC_AVAILABLE = False
except ImportError:
    METADATA_AVAILABLE = False
    print("Warning: Metadata libraries not available. Install piexif, geopy")

# Cost tracking imports (optional - gracefully degrade if not available)
try:
    from cost_roi_calculator import CostROICalculator, CostTracker
    COST_TRACKING_AVAILABLE = True
except ImportError:
    COST_TRACKING_AVAILABLE = False
    # Provide stub implementations for graceful degradation
    class CostTracker:
        """Stub CostTracker when cost tracking is not available."""
        def __init__(self, *args, **kwargs):
            pass
        def __enter__(self):
            return self
        def __exit__(self, *args):
            return False

# Error tracking imports (optional - gracefully degrade if not available)
try:
    from error_tracking import (
        init_sentry,
        capture_error,
        track_operation,
        track_error,
        FileProcessingErrorTracker,
        ErrorLevel
    )
    ERROR_TRACKING_AVAILABLE = True
except ImportError:
    ERROR_TRACKING_AVAILABLE = False
    # Stub implementations
    def init_sentry(*args, **kwargs): return False
    def capture_error(*args, **kwargs): pass
    def track_operation(*args, **kwargs):
        from contextlib import nullcontext
        return nullcontext()
    def track_error(*args, **kwargs):
        def decorator(func): return func
        return decorator
    class FileProcessingErrorTracker:
        def __init__(self): pass
        def track_file(self, *args, **kwargs):
            from contextlib import nullcontext
            return nullcontext()
        def print_summary(self): pass
        def get_stats(self): return {}
    class ErrorLevel:
        FATAL = 'fatal'
        ERROR = 'error'
        WARNING = 'warning'
        INFO = 'info'
        DEBUG = 'debug'

# CLIP enhancement constants for weak image classification
try:
    from shared.constants import (
        CLIP_CATEGORY_PROMPTS,
        CLIP_CONTENT_LABELS,
        CLIP_LABEL_TO_ORGANIZER,
        CLIP_ENHANCE_THRESHOLD,
        CLIP_ENHANCE_HIGH_THRESHOLD,
    )
    ENHANCED_CLIP_AVAILABLE = True
except ImportError:
    ENHANCED_CLIP_AVAILABLE = False


class ContentClassifier:
    """Classifies document content into categories."""

    def __init__(self):
        """Initialize classifier with keyword patterns."""
        # Company name patterns
        self.company_patterns = [
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+LLC\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+L\.L\.C\.\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Inc\.?\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Incorporated\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Corp\.?\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Corporation\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Company\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Co\.\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Ltd\.?\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+Limited\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+LLP\b',
            r'\b([A-Z][A-Za-z0-9\s&\-\.]{2,50})\s+L\.L\.P\.\b',
        ]

        # People name patterns - look for common name patterns
        self.people_patterns = [
            # ALL-CAPS names at start of resume (common in templates)
            # Matches: "ISABEL BUDENZ\nLLM" or "JOHN DOE\nSoftware Engineer"
            r'^([A-Z]{2,})\s+([A-Z]{2,})\s*\n',
            # ALL-CAPS name followed by title/degree
            r'\b([A-Z]{2,})\s+([A-Z]{2,})\s*\n\s*(?:LLM|MBA|PhD|MD|JD|CPA|Software|Engineer|Manager|Director|Analyst)',
            # Name with document type indicators
            r'\b([A-Z][a-z]+)\s+([A-Z][a-z]+)\s+(?:Resume|CV|Cover Letter)\b',
            r'\b([A-Z][a-z]+)\s+([A-Z][a-z]+)\s+(?:Portfolio|Biography|Bio)\b',
            # Field labels followed by names
            r'\b(?:Name|Contact|From|To|Attn|Author|Client|Patient|Student):\s+([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)\b',
            # Email signatures (name before email)
            r'\b([A-Z][a-z]+\s+[A-Z][a-z]+)\s+<[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}>',
            # Name in "Prepared by/for" statements
            r'\b(?:Prepared|Written|Submitted|Signed)\s+(?:by|for):\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\b',
            # Name followed by credentials (MD, PhD, Esq, etc.)
            r'\b([A-Z][a-z]+\s+[A-Z][a-z]+),?\s+(?:MD|PhD|Esq|DDS|CPA|MBA|JD|RN)\b',
            # Mr./Mrs./Ms./Dr. followed by name
            r'\b(?:Mr|Mrs|Ms|Dr|Prof)\.?\s+([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)\b',
            # Name in meeting notes format
            r'\b(?:Attendee|Participant|Speaker|Presenter):\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\b',
        ]

        self.patterns = {
            'legal': {
                'keywords': [
                    'contract', 'agreement', 'terms', 'conditions', 'legal', 'attorney',
                    'law', 'litigation', 'plaintiff', 'defendant', 'court', 'settlement',
                    'lease', 'deed', 'will', 'testament', 'power of attorney', 'notary',
                    'amendment', 'exhibit', 'whereas', 'party', 'parties', 'executed',
                    'operating agreement', 'llc', 'corporation', 'bylaws', 'articles'
                ],
                'subcategories': {
                    'contracts': ['contract', 'agreement', 'terms', 'subscription', 'saas'],
                    'real_estate': ['lease', 'deed', 'property', 'real estate', 'mortgage'],
                    'corporate': ['llc', 'corporation', 'operating agreement', 'bylaws', 'articles', 'formation'],
                    'other': []
                }
            },
            'financial': {
                'keywords': [
                    'invoice', 'receipt', 'tax', 'irs', 'payment', 'bill', 'billing',
                    'statement', 'account', 'balance', 'transaction', 'credit', 'debit',
                    'bank', 'finance', 'loan', 'interest', '1098', '1099', 'w-2', 'w2',
                    'federal', 'state return', 'refund', 'revenue', 'expense', 'budget',
                    'investment', 'portfolio', 'ein', 'employer identification'
                ],
                'subcategories': {
                    'tax': ['tax', 'irs', '1098', '1099', 'w-2', 'w2', 'federal', 'state return'],
                    'invoices': ['invoice', 'bill', 'billing', 'payment'],
                    'statements': ['statement', 'account', 'balance', 'transaction'],
                    'other': []
                }
            },
            'business': {
                'keywords': [
                    'proposal', 'pitch', 'business plan', 'strategy', 'marketing',
                    'presentation', 'deck', 'startup', 'company', 'venture', 'investor',
                    'growth', 'revenue model', 'unit economics', 'expansion', 'rfp',
                    'guidelines', 'program', 'service package', 'pricing', 'client',
                    'customer', 'vendor', 'supplier', 'partner', 'contacts', 'crm',
                    'hiring', 'job posting', 'meeting', 'standup', 'minutes'
                ],
                'subcategories': {
                    'planning': ['business plan', 'strategy', 'expansion', 'growth', 'project'],
                    'marketing': ['marketing', 'pricing', 'service package', 'pitch', 'deck'],
                    'proposals': ['proposal', 'rfp', 'guidelines'],
                    'crm': ['crm', 'contacts', 'microlender', 'customer'],
                    'hr': ['hiring', 'job posting', 'team roster', 'application', 'linkedin'],
                    'meeting_notes': ['meeting', 'standup', 'minutes', 'agenda', 'retrospective'],
                    'clients': ['client', 'llc', 'inc', 'corp', 'company'],  # Legacy
                    'other': []
                }
            },
            'personal': {
                'keywords': [
                    'resume', 'cv', 'cover letter', 'curriculum vitae', 'employment',
                    'personal', 'identification', 'passport', 'driver license', 'ssn',
                    'birth certificate', 'marriage', 'divorce', 'diploma', 'transcript',
                    'reference', 'recommendation'
                ],
                'subcategories': {
                    'employment': ['resume', 'cv', 'cover letter', 'employment', 'reference'],
                    'identification': ['passport', 'driver license', 'ssn', 'id'],
                    'certificates': ['birth certificate', 'marriage', 'divorce', 'diploma'],
                    'other': []
                }
            },
            'medical': {
                'keywords': [
                    'medical', 'health', 'doctor', 'patient', 'prescription', 'diagnosis',
                    'treatment', 'hospital', 'clinic', 'insurance claim', 'hipaa',
                    'vaccination', 'immunization', 'lab results', 'pharmacy'
                ],
                'subcategories': {
                    'records': ['medical record', 'patient', 'diagnosis', 'treatment'],
                    'insurance': ['insurance', 'claim', 'coverage'],
                    'prescriptions': ['prescription', 'pharmacy', 'medication'],
                    'other': []
                }
            },
            'property': {
                'keywords': [
                    'property management', 'tenant', 'landlord', 'rent', 'rental',
                    'maintenance', 'repair', 'inspection', 'utilities', 'hoa'
                ],
                'subcategories': {
                    'leases': ['lease', 'tenant', 'landlord', 'rent', 'rental'],
                    'maintenance': ['maintenance', 'repair', 'inspection'],
                    'other': []
                }
            },
            'education': {
                'keywords': [
                    'course', 'syllabus', 'lecture', 'assignment', 'homework', 'exam',
                    'grade', 'transcript', 'diploma', 'degree', 'certificate', 'university',
                    'college', 'school', 'research paper', 'thesis', 'dissertation'
                ],
                'subcategories': {
                    'coursework': ['course', 'syllabus', 'lecture', 'assignment'],
                    'research': ['research', 'paper', 'thesis', 'dissertation'],
                    'records': ['transcript', 'diploma', 'degree', 'certificate'],
                    'other': []
                }
            },
            'technical': {
                'keywords': [
                    'code', 'software', 'development', 'programming', 'api', 'database',
                    'documentation', 'technical', 'specification', 'architecture', 'design',
                    'system', 'infrastructure', 'deployment', 'configuration'
                ],
                'subcategories': {
                    'documentation': ['documentation', 'spec', 'specification', 'readme'],
                    'architecture': ['architecture', 'design', 'system', 'infrastructure'],
                    'other': []
                }
            },
            'creative': {
                'keywords': [
                    'design', 'graphic', 'illustration', 'artwork', 'photo', 'image',
                    'screenshot', 'mockup', 'prototype', 'wireframe', 'brand', 'logo'
                ],
                'subcategories': {
                    'design': ['design', 'mockup', 'wireframe', 'prototype'],
                    'branding': ['brand', 'logo', 'identity'],
                    'photos': ['photo', 'photography', 'image'],
                    'other': []
                }
            }
        }

    def extract_company_names(self, text: str) -> List[str]:
        """
        Extract company names from text using regex patterns.

        Returns:
            List of detected company names
        """
        companies = []
        for pattern in self.company_patterns:
            matches = re.findall(pattern, text)
            companies.extend(matches)

        # Remove duplicates and clean up
        unique_companies = []
        seen = set()
        for company in companies:
            # Clean up whitespace
            clean = ' '.join(company.split())
            # Skip if too short or already seen
            if len(clean) > 2 and clean.lower() not in seen:
                seen.add(clean.lower())
                unique_companies.append(clean)

        return unique_companies

    def _collapse_spaced_text(self, text: str) -> str:
        """
        Collapse spaced-out text like "I S A B E L  B U D E N Z" to "ISABEL BUDENZ".
        Common in stylized resume/CV templates.
        """
        # Pattern: single letters separated by spaces (at least 3 in a row)
        # Match sequences like "I S A B E L" (single chars with single spaces)
        def collapse_match(match):
            spaced = match.group(0)
            # Remove single spaces between single characters
            collapsed = re.sub(r'(?<=\b[A-Z]) (?=[A-Z]\b)', '', spaced)
            return collapsed

        # Find sequences of spaced single uppercase letters
        # Pattern matches: capital letter, space, capital letter (repeated)
        result = re.sub(r'\b([A-Z] ){2,}[A-Z]\b', collapse_match, text)
        return result

    def extract_people_names(self, text: str) -> List[str]:
        """
        Extract people names from text using regex patterns.

        Returns:
            List of detected people names
        """
        # Preprocess: collapse spaced-out text (common in stylized resumes)
        text = self._collapse_spaced_text(text)

        people = []
        for pattern in self.people_patterns:
            matches = re.findall(pattern, text)
            # Pattern can return tuples (first, last) or single strings
            for match in matches:
                if isinstance(match, tuple):
                    # Join tuple elements (e.g., first name + last name)
                    full_name = ' '.join([m for m in match if m])
                else:
                    full_name = match
                people.append(full_name)

        # Remove duplicates and clean up
        unique_people = []
        seen = set()
        for person in people:
            # Clean up whitespace
            clean = ' '.join(person.split())
            # Convert ALL-CAPS to Title Case (common in resume headers)
            if clean.isupper():
                clean = clean.title()
            # Skip if too short or already seen
            if len(clean) > 2 and clean.lower() not in seen:
                seen.add(clean.lower())
                unique_people.append(clean)

        return unique_people

    def extract_person_company_relationships(self, text: str) -> Dict[str, str]:
        """
        Extract relationships between people and companies from text.
        Uses Schema.org-style connections (Person worksFor/memberOf Organization).

        Returns:
            Dictionary mapping person names to company names
        """
        relationships = {}

        # Patterns for person-company relationships
        relationship_patterns = [
            # "John Doe at Company LLC"
            r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s+(?:at|from)\s+([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))',
            # "John Doe, CEO of Company LLC"
            r'([A-Z][a-z]+\s+[A-Z][a-z]+),?\s+(?:CEO|CFO|CTO|COO|President|Director|Manager|Founder)\s+(?:of|at)\s+([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))',
            # "Company LLC - Contact: John Doe"
            r'([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))\s*[-:]\s*(?:Contact|Representative):\s*([A-Z][a-z]+\s+[A-Z][a-z]+)',
            # "John Doe (Company LLC)"
            r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s+\(([A-Z][A-Za-z0-9\s&\-\.]{2,50}(?:\s+LLC|\s+Inc\.?|\s+Corp\.?|\s+Ltd\.?|\s+LLP))\)',
            # Email pattern: john.doe@company.com -> John Doe at Company
            r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s+<[a-zA-Z0-9._%+-]+@([a-zA-Z0-9.-]+)\.[a-zA-Z]{2,}>',
        ]

        for pattern in relationship_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if len(match) == 2:
                    person, company = match
                    # Clean up
                    person_clean = ' '.join(person.split())
                    company_clean = ' '.join(company.split())

                    # For email domains, capitalize company name
                    if '@' in text and '.' in company_clean and len(company_clean.split('.')) >= 2:
                        # This is likely a domain name, extract company name
                        domain_parts = company_clean.split('.')
                        if domain_parts[0].lower() not in ['gmail', 'yahoo', 'hotmail', 'outlook', 'mail']:
                            company_clean = domain_parts[0].capitalize()

                    # Store relationship (person -> company)
                    if len(person_clean) > 2 and len(company_clean) > 2:
                        relationships[person_clean] = company_clean

        return relationships

    def is_valid_company_name(self, name: str) -> bool:
        """
        Check if a string is a valid company name (not a sentence fragment).

        Returns:
            True if valid company name, False if likely a sentence fragment
        """
        if not name:
            return False

        name_lower = name.lower().strip()
        words = name.split()

        # Reject if too long (real company names are usually < 60 chars)
        if len(name) > 60:
            return False

        # Reject if too many words (company names rarely have > 6 words)
        if len(words) > 6:
            return False

        # Sentence fragment indicators - words that start sentences, not companies
        sentence_starters = {
            'neither', 'either', 'total', 'the', 'a', 'an', 'if', 'when',
            'where', 'while', 'although', 'because', 'since', 'unless',
            'however', 'therefore', 'moreover', 'furthermore', 'additionally',
            'please', 'note', 'see', 'refer', 'click', 'visit', 'contact',
            'for', 'with', 'from', 'into', 'about', 'above', 'below',
            'between', 'under', 'over', 'after', 'before', 'during',
            'this', 'that', 'these', 'those', 'which', 'what', 'who',
            'all', 'any', 'each', 'every', 'both', 'few', 'many', 'most',
            'other', 'some', 'such', 'no', 'not', 'only', 'own', 'same',
            'output', 'input', 'return', 'returns', 'required', 'optional',
        }

        # Check first word
        if words and words[0].lower() in sentence_starters:
            return False

        # Sentence patterns - these indicate full sentences, not company names
        sentence_patterns = [
            r'\b(?:is|are|was|were|be|been|being)\b',  # Verbs
            r'\b(?:to|of|in|on|at|by)\s+(?:the|a|an)\b',  # Preposition + article
            r'\b(?:you|your|we|our|they|their|it|its)\b',  # Pronouns
            r'\b(?:can|could|will|would|shall|should|may|might|must)\b',  # Modal verbs
            r'\b(?:and|or|but|nor|yet|so)\s+\w+\s+\w+',  # Conjunction + multiple words
        ]

        for pattern in sentence_patterns:
            if re.search(pattern, name_lower):
                return False

        # Check for specific problematic patterns
        problematic_phrases = [
            'the name of', 'in usd', 'total in', 'output only',
            'required for', 'agreement between', 'agreement of',
            'certificate of', 'description of', 'operating agreement',
            'license this', 'http rule', 'member-managed',
            'need some', 'print out', 'user provided',
            'ceo of', 'cfo of', 'cto of', 'coo of',  # Title patterns
            'president of', 'director of', 'manager of',
            'taxpayer number', 'tax id', 'ein number',  # Tax/ID patterns
            'student award', 'professional access',  # Award patterns
            'proprietor general', 'general partnership',  # Legal entity types
            'personal workload', 'workload and',  # Incomplete phrases
            'data usage agreement', 'service agreement',
            'contributions on behalf', 'on behalf of',
        ]

        for phrase in problematic_phrases:
            if phrase in name_lower:
                return False

        # Reject names ending with conjunctions (incomplete phrases)
        if words and words[-1].lower() in {'and', 'or', 'but', 'nor', 'yet', 'so', 'the', 'a', 'an', 'of', 'to', 'in', 'on', 'at', 'by'}:
            return False

        # Reject names starting with titles followed by "of"
        if len(words) >= 3 and words[1].lower() == 'of':
            title_words = {'ceo', 'cfo', 'cto', 'coo', 'president', 'director', 'manager', 'chairman', 'founder'}
            if words[0].lower() in title_words:
                return False

        return True

    def normalize_company_name(self, company_name: str) -> str:
        """
        Normalize company name by extracting actual company from common patterns.

        Handles patterns like:
        - "Copyright 2024 Google" -> "Google"
        - "© 2020 Microsoft Corporation" -> "Microsoft"
        - "(c) 2019-2024 Apple Inc" -> "Apple"
        - "Copyright (C) 2023 Amazon" -> "Amazon"
        - "Google LLC" -> "Google"
        - "Apple Inc." -> "Apple"

        Returns:
            Normalized company name, or None if invalid
        """
        if not company_name:
            return company_name

        # Patterns to extract company name from copyright notices
        copyright_patterns = [
            # "Copyright 2024 Google" or "Copyright (C) 2024 Google"
            r'(?:copyright|©|\(c\))\s*(?:\(c\))?\s*(?:\d{4}(?:\s*[-–—]\s*\d{4})?)\s+(.+)',
            # "2024 Google" (just year followed by company)
            r'^\d{4}(?:\s*[-–—]\s*\d{4})?\s+([A-Z][A-Za-z0-9\s&\-\.]+)$',
            # "(c) Google 2024" (company before year)
            r'(?:copyright|©|\(c\))\s+([A-Z][A-Za-z0-9\s&\-\.]+?)\s+\d{4}',
            # "Copyright Google" or "© Google" (without year)
            r'^(?:copyright|©|\(c\))\s+([A-Za-z][A-Za-z0-9\s&\-\.]+)$',
        ]

        name_lower = company_name.lower().strip()
        result = company_name

        # Check if this looks like a copyright notice
        if any(indicator in name_lower for indicator in ['copyright', '©', '(c)']):
            for pattern in copyright_patterns:
                match = re.search(pattern, company_name, re.IGNORECASE)
                if match:
                    extracted = match.group(1).strip()
                    # Clean up trailing punctuation
                    extracted = re.sub(r'[.,;:]+$', '', extracted).strip()
                    if extracted and len(extracted) >= 2:
                        result = extracted
                        break

        # Check for year prefix pattern (e.g., "2024 Google")
        if result == company_name:
            year_prefix_match = re.match(r'^(\d{4}(?:\s*[-–—]\s*\d{4})?)\s+(.+)$', company_name)
            if year_prefix_match:
                extracted = year_prefix_match.group(2).strip()
                if extracted and len(extracted) >= 2:
                    result = extracted

        # Strip legal suffixes to consolidate company variants
        # Order matters: check longer suffixes first
        legal_suffixes = [
            # Full words with variations
            r'\s+Incorporated$',
            r'\s+Corporation$',
            r'\s+Limited$',
            r'\s+Company$',
            # Abbreviations with optional period
            r'\s+L\.L\.C\.$',
            r'\s+L\.L\.P\.$',
            r'\s+LLC\.?$',
            r'\s+LLP\.?$',
            r'\s+Inc\.?$',
            r'\s+Corp\.?$',
            r'\s+Ltd\.?$',
            r'\s+Co\.?$',
            # Other common suffixes
            r'\s+PLC\.?$',
            r'\s+LP\.?$',
            r'\s+SA$',
            r'\s+GmbH$',
            r'\s+AG$',
        ]

        for suffix_pattern in legal_suffixes:
            result = re.sub(suffix_pattern, '', result, flags=re.IGNORECASE).strip()

        return result

    def sanitize_company_name(self, company_name: str) -> Optional[str]:
        """
        Sanitize company name for use in folder names.

        Returns:
            Sanitized folder name, or None if the name is invalid (sentence fragment)
        """
        # First normalize the company name (extract from copyright patterns, etc.)
        normalized = self.normalize_company_name(company_name)

        # Validate that this is a real company name, not a sentence fragment
        if not self.is_valid_company_name(normalized):
            return None

        # Remove special characters that aren't allowed in folder names
        sanitized = re.sub(r'[<>:"/\\|?*]', '', normalized)
        # Replace multiple spaces with single space
        sanitized = ' '.join(sanitized.split())
        # Limit length
        if len(sanitized) > 50:
            sanitized = sanitized[:50].strip()
        return sanitized if sanitized else None

    def classify_content(self, text: str, filename: str = "") -> Tuple[str, str, Optional[str], List[str]]:
        """
        Classify content based on extracted text.
        Uses Schema.org person-company relationships to improve categorization.

        Returns:
            Tuple of (category, subcategory, company_name, people_names)
        """
        if not text:
            return ('uncategorized', 'other', None, [])

        text_lower = text.lower()
        filename_lower = filename.lower()
        combined = f"{text_lower} {filename_lower}"

        # Check for known companies in text (canonical name mapping)
        known_text_companies = {
            'capital city village': ('organization', 'property_management', 'Capital City Village'),
            'leora home health': ('organization', 'healthcare', 'Leora Home Health'),
            'integrity studio': ('organization', 'vendors', 'Integrity Studio'),
            'inspired movement': ('organization', 'vendors', 'Inspired Movement'),
            'new beginnings child development': ('organization', 'vendors', 'New Beginnings Child Development Center'),
            'zouk': ('zouk', 'events', None),
        }
        for phrase, (cat, subcat, canonical_name) in known_text_companies.items():
            if phrase in text_lower:
                return (cat, subcat, canonical_name, self.extract_people_names(text))

        # Extract company names and people names
        company_names = self.extract_company_names(text)
        primary_company = company_names[0] if company_names else None

        people_names = self.extract_people_names(text)

        # Extract person-company relationships (Schema.org connections)
        person_company_relationships = self.extract_person_company_relationships(text)

        # Prioritize company from person-company relationships over direct extraction
        # Relationships tend to be more accurate as they include context
        if person_company_relationships:
            # Get the first relationship's company
            relationship_company = next(iter(person_company_relationships.values()))

            # Check if relationship company has proper legal suffix
            has_legal_suffix = any(relationship_company.endswith(suffix)
                                  for suffix in ['LLC', 'Inc.', 'Inc', 'Corp.', 'Corp',
                                                'Ltd.', 'Ltd', 'LLP', 'L.L.C.', 'L.L.P.'])

            # Prefer relationship company if it has legal suffix or we don't have a primary company
            if has_legal_suffix or not primary_company:
                primary_company = relationship_company
            # Or if the relationship company is much cleaner (shorter and no weird prefixes)
            elif primary_company and 'CEO' not in relationship_company and 'at' not in relationship_company:
                if len(relationship_company) < len(primary_company) * 0.8:
                    primary_company = relationship_company

        # Fallback: If we found people but no company, check relationships again
        if people_names and not primary_company and person_company_relationships:
            # Try to find a company for the first person mentioned
            for person in people_names:
                if person in person_company_relationships:
                    primary_company = person_company_relationships[person]
                    break

        # Score each category
        scores = defaultdict(int)
        category_subcats = {}

        for category, data in self.patterns.items():
            for keyword in data['keywords']:
                count = combined.count(keyword.lower())
                if count > 0:
                    scores[category] += count

                    # Track which subcategory keywords matched
                    for subcat, subcat_keywords in data['subcategories'].items():
                        if any(sk.lower() in combined for sk in subcat_keywords):
                            if category not in category_subcats:
                                category_subcats[category] = defaultdict(int)
                            category_subcats[category][subcat] += count

        if not scores:
            return ('uncategorized', 'other', primary_company, people_names)

        # Get category with highest score
        best_category = max(scores.items(), key=lambda x: x[1])[0]

        # Get subcategory with highest score for this category
        if best_category in category_subcats:
            subcat_scores = category_subcats[best_category]
            if subcat_scores:
                best_subcategory = max(subcat_scores.items(), key=lambda x: x[1])[0]
            else:
                best_subcategory = 'other'
        else:
            best_subcategory = 'other'

        # If we detected a company (either directly or via person relationship)
        # and it's business-related, use clients subcategory
        if primary_company and best_category == 'business':
            best_subcategory = 'clients'

        return (best_category, best_subcategory, primary_company, people_names)


class ImageMetadataParser:
    """Parses image metadata including EXIF, GPS, and timestamps."""

    def __init__(self, cost_calculator: 'CostROICalculator' = None):
        """
        Initialize the metadata parser.

        Args:
            cost_calculator: Optional cost calculator for tracking usage costs
        """
        self.metadata_available = METADATA_AVAILABLE
        self.geocoder = None
        self.cost_calculator = cost_calculator

        if self.metadata_available:
            try:
                # Initialize geocoder with a user agent
                self.geocoder = Nominatim(user_agent="file_organizer_v1.0", timeout=5)
            except Exception as e:
                print(f"Warning: Could not initialize geocoder: {e}")
                self.geocoder = None

    def extract_exif_data(self, image_path: Path) -> Dict[str, Any]:
        """
        Extract EXIF data from an image.

        Returns:
            Dictionary with EXIF data
        """
        if not self.metadata_available:
            return {}

        try:
            image = Image.open(image_path)
            exif_data = {}

            # Get EXIF data
            exif = image._getexif()
            if exif:
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    exif_data[tag] = value

            return exif_data

        except Exception as e:
            print(f"  EXIF extraction error: {e}")
            return {}

    def extract_datetime(self, image_path: Path) -> Optional[datetime]:
        """
        Extract the datetime when the photo was taken.

        Returns:
            datetime object or None
        """
        exif_data = self.extract_exif_data(image_path)

        if not exif_data:
            return None

        # Try different datetime tags
        datetime_tags = ['DateTimeOriginal', 'DateTimeDigitized', 'DateTime']

        for tag in datetime_tags:
            if tag in exif_data:
                try:
                    # Parse datetime string (format: "2023:11:26 14:30:00")
                    dt_str = str(exif_data[tag])
                    dt = datetime.strptime(dt_str, "%Y:%m:%d %H:%M:%S")
                    return dt
                except (ValueError, TypeError):
                    continue

        return None

    def extract_gps_coordinates(self, image_path: Path) -> Optional[Tuple[float, float]]:
        """
        Extract GPS coordinates from image EXIF data.

        Returns:
            Tuple of (latitude, longitude) or None
        """
        if not self.metadata_available:
            return None

        try:
            image = Image.open(image_path)
            exif = image._getexif()

            if not exif:
                return None

            # Get GPS info
            gps_info = {}
            for tag_id, value in exif.items():
                tag = TAGS.get(tag_id, tag_id)
                if tag == 'GPSInfo':
                    for gps_tag_id in value:
                        gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                        gps_info[gps_tag] = value[gps_tag_id]

            if not gps_info:
                return None

            # Convert to decimal degrees
            lat = self._convert_to_degrees(gps_info.get('GPSLatitude'))
            lon = self._convert_to_degrees(gps_info.get('GPSLongitude'))

            if lat is None or lon is None:
                return None

            # Adjust for hemisphere
            if gps_info.get('GPSLatitudeRef') == 'S':
                lat = -lat
            if gps_info.get('GPSLongitudeRef') == 'W':
                lon = -lon

            return (lat, lon)

        except Exception as e:
            print(f"  GPS extraction error: {e}")
            return None

    def _convert_to_degrees(self, value) -> Optional[float]:
        """
        Convert GPS coordinates to degrees.

        Args:
            value: GPS coordinate in format ((deg, 1), (min, 1), (sec, 1))

        Returns:
            Decimal degrees or None
        """
        if not value:
            return None

        try:
            d = float(value[0][0]) / float(value[0][1])
            m = float(value[1][0]) / float(value[1][1])
            s = float(value[2][0]) / float(value[2][1])

            return d + (m / 60.0) + (s / 3600.0)
        except (IndexError, TypeError, ZeroDivisionError):
            return None

    def get_location_name(self, coordinates: Tuple[float, float]) -> Optional[str]:
        """
        Get location name from GPS coordinates using reverse geocoding.

        Args:
            coordinates: Tuple of (latitude, longitude)

        Returns:
            Location name (city, state, country) or None
        """
        if not self.geocoder:
            return None

        with CostTracker(self.cost_calculator, 'nominatim_geocoding') if self.cost_calculator else nullcontext():
            try:
                lat, lon = coordinates
                location = self.geocoder.reverse(f"{lat}, {lon}", exactly_one=True)

                if location and location.raw.get('address'):
                    address = location.raw['address']

                    # Try to get city, state, country
                    parts = []

                    # City
                    city = address.get('city') or address.get('town') or address.get('village')
                    if city:
                        parts.append(city)

                    # State/Region
                    state = address.get('state') or address.get('region')
                    if state:
                        parts.append(state)

                    # Country
                    country = address.get('country')
                    if country:
                        parts.append(country)

                    if parts:
                        return ', '.join(parts)

            except (GeocoderTimedOut, GeocoderServiceError) as e:
                print(f"  Geocoding error: {e}")
            except Exception as e:
                print(f"  Location lookup error: {e}")

            return None

    def get_metadata_summary(self, image_path: Path) -> Dict[str, Any]:
        """
        Get a summary of image metadata.

        Returns:
            Dictionary with datetime, GPS coordinates, and location
        """
        summary = {
            'datetime': None,
            'gps_coordinates': None,
            'location_name': None,
            'year': None,
            'month': None,
            'date_str': None
        }

        # Extract datetime
        dt = self.extract_datetime(image_path)
        if dt:
            summary['datetime'] = dt
            summary['year'] = dt.year
            summary['month'] = dt.month
            summary['date_str'] = dt.strftime("%Y-%m")

        # Extract GPS
        coords = self.extract_gps_coordinates(image_path)
        if coords:
            summary['gps_coordinates'] = coords
            # Get location name (with rate limiting consideration)
            location = self.get_location_name(coords)
            if location:
                summary['location_name'] = location

        return summary


class ImageContentAnalyzer:
    """Analyzes image content using computer vision."""

    def __init__(self, cost_calculator: 'CostROICalculator' = None):
        """
        Initialize the image content analyzer.

        Args:
            cost_calculator: Optional cost calculator for tracking usage costs
        """
        self.vision_available = VISION_AVAILABLE
        self.model = None
        self.processor = None
        self.face_cascade = None
        self.cost_calculator = cost_calculator

        if self.vision_available:
            try:
                print("Loading CLIP model for image analysis...")
                self.model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
                self.processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")

                # Load OpenCV face detection
                cascade_path = cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
                self.face_cascade = cv2.CascadeClassifier(cascade_path)

                print("✓ CLIP model loaded successfully")
            except Exception as e:
                print(f"Warning: Could not load CLIP model: {e}")
                self.vision_available = False

    def detect_people(self, image_path: Path) -> bool:
        """
        Detect if there are people in the image using face detection.

        Returns:
            True if people detected, False otherwise
        """
        if not self.vision_available or self.face_cascade is None:
            return False

        with CostTracker(self.cost_calculator, 'face_detection') if self.cost_calculator else nullcontext():
            try:
                # Read image
                img = cv2.imread(str(image_path))
                if img is None:
                    return False

                # Convert to grayscale for face detection
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

                # Detect faces
                faces = self.face_cascade.detectMultiScale(
                    gray,
                    scaleFactor=1.1,
                    minNeighbors=5,
                    minSize=(30, 30)
                )

                return len(faces) > 0

            except Exception as e:
                print(f"  Face detection error: {e}")
                return False

    def classify_image_content(self, image_path: Path) -> Dict[str, float]:
        """
        Classify image content using CLIP zero-shot classification.

        Returns:
            Dictionary of category -> confidence score
        """
        if not self.vision_available or self.model is None:
            return {}

        with CostTracker(self.cost_calculator, 'clip_vision') if self.cost_calculator else nullcontext():
            try:
                # Open image
                image = Image.open(image_path)

                # Define categories to check
                categories = [
                    "a photo of a home interior room",
                    "a photo of a living room",
                    "a photo of a bedroom",
                    "a photo of a kitchen",
                    "a photo of a bathroom",
                    "a photo of furniture",
                    "a photo of a house exterior",
                    "a photo of people",
                    "a screenshot",
                    "a photo of outdoors",
                    "a photo of nature"
                ]

                # Prepare inputs
                inputs = self.processor(
                    text=categories,
                    images=image,
                    return_tensors="pt",
                    padding=True
                )

                # Get predictions
                with torch.no_grad():
                    outputs = self.model(**inputs)
                    logits_per_image = outputs.logits_per_image
                    probs = logits_per_image.softmax(dim=1)

                # Convert to dictionary
                scores = {}
                for i, category in enumerate(categories):
                    scores[category] = float(probs[0][i])

                return scores

            except Exception as e:
                print(f"  Image classification error: {e}")
            return {}

    def is_home_interior_no_people(self, image_path: Path) -> Tuple[bool, Dict[str, float]]:
        """
        Check if image is a home interior without people.

        Returns:
            Tuple of (is_interior_no_people, classification_scores)
        """
        if not self.vision_available:
            return (False, {})

        # Classify image content
        scores = self.classify_image_content(image_path)

        if not scores:
            return (False, {})

        # Check for home interior indicators
        interior_score = max(
            scores.get("a photo of a home interior room", 0),
            scores.get("a photo of a living room", 0),
            scores.get("a photo of a bedroom", 0),
            scores.get("a photo of a kitchen", 0),
            scores.get("a photo of a bathroom", 0),
            scores.get("a photo of furniture", 0)
        )

        # Check for people
        people_score = scores.get("a photo of people", 0)
        has_faces = self.detect_people(image_path)

        # Determine if it's an interior without people
        is_interior = interior_score > 0.3  # 30% confidence threshold
        has_people = people_score > 0.2 or has_faces  # 20% confidence or face detection

        return (is_interior and not has_people, scores)

    def has_people_in_photo(self, image_path: Path) -> Tuple[bool, Dict[str, float]]:
        """
        Check if image contains people (for social photos).

        Returns:
            Tuple of (has_people, classification_scores)
        """
        if not self.vision_available:
            return (False, {})

        # Classify image content
        scores = self.classify_image_content(image_path)

        if not scores:
            return (False, {})

        # Check for people indicators
        people_score = scores.get("a photo of people", 0)
        has_faces = self.detect_people(image_path)

        # Check that it's NOT a screenshot
        screenshot_score = scores.get("a screenshot", 0)
        is_screenshot = screenshot_score > 0.4  # High threshold for screenshots

        # Determine if photo has people (and is not a screenshot)
        has_people = (people_score > 0.15 or has_faces) and not is_screenshot

        return (has_people, scores)


class ContentBasedFileOrganizer:
    """Organize files based on content analysis using OCR."""

    def __init__(
        self,
        base_path: str = None,
        organize_by_date: bool = False,
        organize_by_location: bool = False,
        enable_cost_tracking: bool = True,
        db_path: str = 'results/file_organization.db'
    ):
        """
        Initialize the organizer.

        Args:
            base_path: Base path for organized files
            organize_by_date: If True, organize photos by date (Photos/2023/11/)
            organize_by_location: If True, organize photos by location when GPS data available
            enable_cost_tracking: If True, track costs and ROI for all features
            db_path: Path to SQLite database for persistent storage
        """
        self.base_path = Path(base_path or "~/Documents").expanduser()

        # Initialize cost tracking if available and enabled
        self.cost_calculator = None
        if enable_cost_tracking and COST_TRACKING_AVAILABLE:
            self.cost_calculator = CostROICalculator()
            print("✓ Cost tracking enabled")

        # Initialize graph store for persistent storage with canonical IDs
        self.graph_store = None
        if GRAPH_STORE_AVAILABLE and db_path:
            self.graph_store = GraphStore(db_path=db_path)
            print(f"✓ Graph store enabled ({db_path})")

        self.enricher = MetadataEnricher()
        self.validator = SchemaValidator()
        self.registry = SchemaRegistry()
        self.classifier = ContentClassifier()
        self.image_analyzer = ImageContentAnalyzer(cost_calculator=self.cost_calculator)
        self.metadata_parser = ImageMetadataParser(cost_calculator=self.cost_calculator)
        self.stats = defaultdict(int)
        self.ocr_available = OCR_AVAILABLE
        self.organize_by_date = organize_by_date
        self.organize_by_location = organize_by_location

        # Filepath-based classification (checked FIRST before content analysis)
        self.filepath_patterns = {
            # Log files
            '.log': 'Technical/Logs',
            '.log.gz': 'Technical/Logs',
            '.out': 'Technical/Logs',

            # Python
            '.py': 'Technical/Python',
            '.pyc': 'Technical/Python/Compiled',
            '.pyw': 'Technical/Python',
            '.pyx': 'Technical/Python',
            '.pyd': 'Technical/Python',

            # JavaScript/TypeScript
            '.js': 'Technical/JavaScript',
            '.jsx': 'Technical/JavaScript',
            '.mjs': 'Technical/JavaScript',
            '.cjs': 'Technical/JavaScript',
            '.ts': 'Technical/TypeScript',
            '.tsx': 'Technical/TypeScript',

            # Web
            '.html': 'Technical/Web',
            '.htm': 'Technical/Web',
            '.css': 'Technical/Web',
            '.scss': 'Technical/Web',
            '.sass': 'Technical/Web',
            '.less': 'Technical/Web',

            # Shell scripts
            '.sh': 'Technical/Shell',
            '.bash': 'Technical/Shell',
            '.zsh': 'Technical/Shell',
            '.fish': 'Technical/Shell',

            # Config files
            '.json': 'Technical/Config',
            '.yaml': 'Technical/Config',
            '.yml': 'Technical/Config',
            '.toml': 'Technical/Config',
            '.ini': 'Technical/Config',
            '.conf': 'Technical/Config',
            '.config': 'Technical/Config',
            '.env': 'Technical/Config',

            # Database
            '.sql': 'Technical/Database',
            '.db': 'Technical/Database',
            '.sqlite': 'Technical/Database',
            '.sqlite3': 'Technical/Database',

            # Java/Kotlin
            '.java': 'Technical/Java',
            '.class': 'Technical/Java/Compiled',
            '.jar': 'Technical/Java/Archives',
            '.kt': 'Technical/Kotlin',
            '.kts': 'Technical/Kotlin',

            # C/C++
            '.c': 'Technical/C',
            '.cpp': 'Technical/C++',
            '.cc': 'Technical/C++',
            '.cxx': 'Technical/C++',
            '.h': 'Technical/C/Headers',
            '.hpp': 'Technical/C++/Headers',

            # Go
            '.go': 'Technical/Go',

            # Rust
            '.rs': 'Technical/Rust',

            # Ruby
            '.rb': 'Technical/Ruby',
            '.rake': 'Technical/Ruby',

            # PHP
            '.php': 'Technical/PHP',

            # Swift
            '.swift': 'Technical/Swift',

            # Markdown and docs
            '.md': 'Technical/Documentation',
            '.markdown': 'Technical/Documentation',
            '.rst': 'Technical/Documentation',
            '.adoc': 'Technical/Documentation',

            # Version control
            '.gitignore': 'Technical/VersionControl',
            '.gitattributes': 'Technical/VersionControl',

            # Build/Package files
            'Makefile': 'Technical/Build',
            'Dockerfile': 'Technical/Build',
            'docker-compose.yml': 'Technical/Build',
            'package.json': 'Technical/Build',
            'package-lock.json': 'Technical/Build',
            'yarn.lock': 'Technical/Build',
            'Cargo.toml': 'Technical/Build',
            'go.mod': 'Technical/Build',
            'requirements.txt': 'Technical/Build',
            'Pipfile': 'Technical/Build',
            'pyproject.toml': 'Technical/Build',
        }

        # Content-based organization structure
        self.category_paths = {
            'legal': {
                'contracts': 'Legal/Contracts',
                'real_estate': 'Legal/RealEstate',
                'corporate': 'Legal/Corporate',
                'other': 'Legal/Other'
            },
            'financial': {
                'tax': 'Financial/Tax',
                'invoices': 'Financial/Invoices',
                'statements': 'Financial/Statements',
                'other': 'Financial/Other'
            },
            'business': {
                'planning': 'Business/Planning',
                'marketing': 'Business/Marketing',
                'proposals': 'Business/Proposals',
                'presentations': 'Business/Presentations',
                'crm': 'Business/CRM',
                'hr': 'Business/HR',
                'meeting_notes': 'Business/MeetingNotes',
                'clients': 'Business/Clients',  # Legacy - prefer crm/hr
                'other': 'Business/Other'
            },
            'personal': {
                'employment': 'Personal/Employment',
                'identification': 'Personal/Identification',
                'certificates': 'Personal/Certificates',
                'other': 'Personal/Other'
            },
            'medical': {
                'records': 'Medical/Records',
                'insurance': 'Medical/Insurance',
                'prescriptions': 'Medical/Prescriptions',
                'other': 'Medical/Other'
            },
            'property': {
                'leases': 'Property/Leases',
                'maintenance': 'Property/Maintenance',
                'other': 'Property/Other'
            },
            'education': {
                'coursework': 'Education/Coursework',
                'research': 'Education/Research',
                'records': 'Education/Records',
                'other': 'Education/Other'
            },
            'technical': {
                'documentation': 'Technical/Documentation',
                'architecture': 'Technical/Architecture',
                'config': 'Technical/Config',
                'data': 'Technical/Data',
                'logs': 'Technical/Logs',
                'web': 'Technical/Web',
                'software_packages': 'Technical/Software_Packages',
                'other': 'Technical/Other'
            },
            'creative': {
                'design': 'Creative/Design',
                'branding': 'Creative/Branding',
                'photos': 'Creative/Photos',
                'other': 'Creative/Other'
            },
            'property_management': 'Property_Management',
            'zouk': {
                'events': 'Zouk/Events',
                'classes': 'Zouk/Classes',
                'other': 'Zouk/Other'
            },
            # Organization: root folder with entity-named subfolders
            # Structure: Organization/{OrgName}/ for most types
            # Exception: Organization/Clients/{OrgName}/ for clients (nested)
            'organization': {
                'clients': 'Organization/Clients',  # Gets nested subfolders
                'vendors': 'Organization',          # Root folder, entity name added dynamically
                'partners': 'Organization',
                'employers': 'Organization',
                'government': 'Organization',
                'healthcare': 'Organization',
                'property_management': 'Organization',
                'financial': 'Organization',
                'educational': 'Organization',
                'nonprofit': 'Organization',
                'meeting_notes': 'Organization',    # Gets Meeting Notes subfolder after company
                'other': 'Organization'
            },
            # Person: root folder with person-named subfolders
            # Structure: Person/{PersonName}/ for all types
            'person': {
                'contacts': 'Person',               # Root folder, person name added dynamically
                'employees': 'Person',
                'clients': 'Person',
                'family': 'Person',
                'references': 'Person',
                'travel': 'Person/Travel',
                'events': 'Person/Events',
                'journal': 'Person/Journal',        # Personal writing, dreams, reflections
                'other': 'Person'
            },
            'game_assets': {
                'audio': 'GameAssets/Audio',
                'music': 'GameAssets/Music',
                'sprites': 'GameAssets/Sprites',
                'textures': 'GameAssets/Textures',
                'fonts': 'GameAssets/Fonts',
                'other': 'GameAssets/Other'
            },
            'fonts': {
                'truetype': 'CreativeWork/Fonts/TrueType',
                'opentype': 'CreativeWork/Fonts/OpenType',
                'web': 'CreativeWork/Fonts/Web',
                'other': 'CreativeWork/Fonts/Other'
            },
            'media': {
                'photos': {
                    'screenshots': {
                        'browser': 'Media/Photos/Screenshots/Browser',
                        'terminal': 'Media/Photos/Screenshots/Terminal',
                        'code': 'Media/Photos/Screenshots/CodeEditors',
                        'docs': 'Media/Photos/Screenshots/Docs',
                        'settings': 'Media/Photos/Screenshots/Settings',
                        'products': 'Media/Photos/Screenshots/Products',
                        'dashboard': 'Media/Photos/Screenshots/Dashboards',
                        'chat': 'Media/Photos/Screenshots/Chat',
                        'other': 'Media/Photos/Screenshots',
                    },
                    'travel': 'Media/Photos/Travel',
                    'portraits': 'Media/Photos/Portraits',
                    'events': 'Media/Photos/Events',
                    'documents': 'Media/Photos/Documents',
                    'social': 'Media/Photos/Social',
                    'chatgpt': 'Media/Photos/ChatGPT',
                    'facebook': 'Media/Photos/Facebook',
                    'logos': 'Media/Photos/Logos',
                    'stock': 'Media/Photos/Stock',
                    'nature': 'Media/Photos/Nature',
                    'lifestyle': 'Media/Photos/Lifestyle',
                    'products': 'Media/Photos/Products',
                    'other': 'Media/Photos/Other'
                },
                'videos': {
                    'recordings': 'Media/Videos/Recordings',
                    'exports': 'Media/Videos/Exports',
                    'screencasts': 'Media/Videos/Screencasts',
                    'other': 'Media/Videos/Other'
                },
                'audio': {
                    'recordings': 'Media/Audio/Recordings',
                    'music': 'Media/Audio/Music',
                    'podcasts': 'Media/Audio/Podcasts',
                    'other': 'Media/Audio/Other'
                },
                'graphics': {
                    'vector': 'Media/Graphics/Vector',
                    'icons': 'Media/Graphics/Icons',
                    'other': 'Media/Graphics/Other'
                },
                'other': 'Media/Other'
            },
            'uncategorized': 'Uncategorized'
        }

        # Game asset detection patterns
        self.game_audio_keywords = [
            'bolt', 'spell', 'magic', 'cast', 'chirp', 'crossbow', 'dagger',
            'sword', 'arrow', 'bow', 'heal', 'potion', 'lightning', 'fire',
            'ice', 'acid', 'poison', 'explosion', 'blast', 'summon', 'dispel',
            'petrification', 'neutralize', 'slow', 'darkness', 'achievement',
            'quest', 'unlock', 'lock', 'door', 'chest', 'coin', 'pickup',
            'attack', 'hit', 'damage', 'death', 'footstep', 'jump', 'land',
            'monster', 'creature', 'enemy', 'boss', 'battle', 'combat',
            'starving', 'hunger', 'thirst', 'eat', 'drink', 'sleep',
            'fiddle', 'lute', 'mandoline', 'glockenspiel', 'instrument',
            'identify', 'greater', 'mental'
        ]

        self.game_music_keywords = [
            'battle', 'boss', 'dungeon', 'castle', 'forest', 'town', 'village',
            'temple', 'ruins', 'cave', 'mountain', 'ocean', 'desert', 'snow',
            'victory', 'defeat', 'theme', 'menu', 'credits', 'intro', 'outro',
            'mysterious', 'dark', 'light', 'epic', 'calm', 'peaceful', 'tension',
            'chaos', 'hope', 'despair', 'triumph', 'march', 'symphony', 'monotony',
            'drakalor', 'altar', 'lawful', 'chaotic', 'neutral', 'alignment',
            'dwarven', 'elven', 'orcish', 'halls', 'abandon', 'corrupting',
            'breeze', 'clockwork', 'knowledge', 'oddisey', 'final', 'welcome'
        ]

        self.game_sprite_keywords = [
            'frame', 'item', 'segment', 'sprite', 'texture', 'tile',
            'leg', 'arm', 'head', 'torso', 'body', 'wing', 'tail',
            'hair', 'face', 'eye', 'mouth', 'hand', 'foot',
            # Character customization sprites
            'beard', 'bling', 'hiero', 'mustache', 'scar', 'tattoo',
            'earring', 'necklace', 'bracelet', 'glasses', 'mask', 'hood',
            'wall', 'floor', 'ceiling', 'door', 'window', 'stairs',
            'tree', 'rock', 'grass', 'water', 'lava', 'cloud',
            'sword', 'shield', 'armor', 'helmet', 'boot', 'glove',
            'potion', 'scroll', 'wand', 'staff', 'ring', 'amulet',
            'coin', 'gem', 'crystal', 'ore', 'metal', 'wood',
            'monster', 'enemy', 'npc', 'character', 'player', 'hero',
            'icon', 'button', 'ui', 'hud', 'menu', 'cursor',
            'particle', 'effect', 'explosion', 'smoke', 'blood',
            'corner', 'edge', 'border', 'container', 'btn', 'talent',
            'bar', 'over', 'down', 'up', 'left', 'right', 'main',
            'bottom', 'top', 'extension', 'descend', 'ascend',
            'mad_carpenter', 'no_more', 'bedroom', 'front', 'back',
            'upper', 'lower', 'middle', 'dead', 'alive', 'sleeping',
            'female', 'male', 'white', 'black', 'red', 'blue', 'green',
            'silver', 'gold', 'bronze', 'iron', 'steel', 'mithril',
            'hills', 'road', 'path', 'bridge', 'gate', 'fence',
            # Additional game environment and UI keywords
            'tentacle', 'shadow', 'altar', 'dungeon', 'throne', 'torch',
            'cloak', 'champion', 'curse', 'warning', 'mouse', 'slider',
            'decal', 'column', 'banner', 'sewer', 'statue', 'pillar',
            'orc', 'dwarf', 'elf', 'hurth', 'helf', 'troll', 'goblin',
            'fire', 'ice', 'sand', 'mount', 'tmount', 'deco', 'entrance',
            'pupils', 'shoulders', 'stunned', 'poisoned', 'blind', 'deaf',
            'slowed', 'levitating', 'hungry', 'strained', 'next', 'prev',
            'groove', 'handle', 'cube', 'psf', 'inventory',
            # Weapons and equipment
            '2h_axe', '2h_hammer', '1h_sword', '1h_axe', 'crossbow', 'longbow',
            'dagger', 'mace', 'flail', 'spear', 'halberd', 'scimitar',
            # Skills and abilities
            'assassins_deed', 'atonement', 'backstab', 'cleave', 'smite',
            'fireball', 'lightning', 'heal', 'buff', 'debuff', 'aura',
            # UI and icons
            'arrow_v', 'arrow_h', 'checkbox', 'radio', 'toggle', 'add',
            # Grayscale/variant markers (common in game assets)
            '_grey', '_gray', '_disabled', '_hover', '_active', '_pressed',
            '_selected', '_normal', '_highlight', '_glow', '_dark', '_light'
        ]

        # Regex patterns for game asset detection (numbered sprites, variants)
        import re
        self.game_sprite_patterns = [
            re.compile(r'^\d+_\d+$'),  # 42_8, 51_3, 16_3 (sprite sheets)
            re.compile(r'^\d+_grey(_\d+)?$', re.IGNORECASE),  # 10_grey, 10_grey_1
            re.compile(r'^\d+_f(_\d+)?$', re.IGNORECASE),  # 283_f, 283_f_1
            re.compile(r'^[a-z]+_\d+$', re.IGNORECASE),  # frame_1, item_42
            re.compile(r'^[a-z]+_[a-z]+_\d+$', re.IGNORECASE),  # assassins_deed_1
            re.compile(r'^\d+h_[a-z]+(_\d+)?$', re.IGNORECASE),  # 2h_axe, 2h_axe_1
            re.compile(r'^[a-z]+_v(_\d+)?$', re.IGNORECASE),  # arrow_v, arrow_v_1
            re.compile(r'^[a-z]+_h(_\d+)?$', re.IGNORECASE),  # arrow_h, arrow_h_1
            re.compile(r'^(head|torso|arm|leg|body|wing|hair)_\w+', re.IGNORECASE),  # body parts
            re.compile(r'^(weapon|armor|item|sprite|frame|tile)\d*_', re.IGNORECASE),  # game prefixes
        ]

        # Game font sprite sheet patterns
        self.game_font_keywords = [
            'broguefont', 'gamefont', 'pixelfont', 'bitfont', 'font_',
            '_font', 'fontsheet', 'font_atlas', 'fontatlas', 'charset',
            'glyphs', 'tilefont', 'asciifont', 'ascii_font'
        ]

    def classify_by_filepath(self, file_path: Path) -> Optional[str]:
        """
        Classify file based on filepath patterns (extension, filename).

        Returns:
            Category path string if matched, None otherwise
        """
        # Check exact filename matches first (e.g., Makefile, Dockerfile)
        filename = file_path.name
        if filename in self.filepath_patterns:
            return self.filepath_patterns[filename]

        # Check file extension
        ext = file_path.suffix.lower()
        if ext in self.filepath_patterns:
            base_path = self.filepath_patterns[ext]

            # Try to extract project name from path
            project_name = self.extract_project_name(file_path)
            if project_name:
                # Add project subdirectory (e.g., Technical/Python/MyProject)
                return f"{base_path}/{project_name}"

            return base_path

        # Check double extensions (e.g., .log.gz)
        if len(file_path.suffixes) >= 2:
            double_ext = ''.join(file_path.suffixes[-2:]).lower()
            if double_ext in self.filepath_patterns:
                return self.filepath_patterns[double_ext]

        return None

    def extract_project_name(self, file_path: Path) -> Optional[str]:
        """
        Extract project name from file path.

        Looks for common project indicators in path:
        - Directory names like 'myproject', 'my-app', etc.
        - Skips common non-project directories

        Returns:
            Project name if found, None otherwise
        """
        skip_dirs = {
            'src', 'lib', 'bin', 'dist', 'build', 'out', 'target',
            'node_modules', 'venv', '.venv', 'env', '__pycache__',
            'scripts', 'tests', 'test', 'docs', 'doc', 'examples',
            'static', 'public', 'assets', 'resources', 'config',
            'home', 'users', 'documents', 'downloads', 'desktop',
            'code', 'projects', 'dev', 'work', 'repos', 'git'
        }

        # Get all parent directories
        parts = file_path.parts

        # Look backwards from the file for a likely project directory
        for i in range(len(parts) - 2, -1, -1):  # Skip the filename itself
            dir_name = parts[i].lower()

            # Skip common non-project directories
            if dir_name in skip_dirs:
                continue

            # Skip hidden directories
            if dir_name.startswith('.'):
                continue

            # Found a likely project directory
            # Return with original case preserved
            return parts[i]

        return None

    def classify_game_asset(self, file_path: Path) -> Optional[Tuple[str, str]]:
        """
        Classify file as a game asset based on filename patterns.

        Returns:
            Tuple of (category, subcategory) or None if not a game asset
        """
        filename = file_path.name.lower()
        stem = file_path.stem.lower()
        ext = file_path.suffix.lower()

        # Remove timestamp suffixes for pattern matching (e.g., _20251120_164506)
        import re
        clean_stem = re.sub(r'_\d{8}_\d{6}$', '', stem)

        # Check for audio files (.wav, .ogg, .mp3)
        if ext in ['.wav', '.ogg', '.mp3', '.flac', '.aac']:
            # Check for game music patterns (usually .ogg files with specific names)
            if ext == '.ogg':
                for keyword in self.game_music_keywords:
                    if keyword in stem:
                        return ('game_assets', 'music')

            # Check for game sound effects
            for keyword in self.game_audio_keywords:
                if keyword in stem:
                    return ('game_assets', 'audio')

        # Check for image files that are game sprites/textures
        if ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tga', '.dds']:
            # Check for game font sprite sheets first
            for keyword in self.game_font_keywords:
                if keyword in stem or keyword in clean_stem:
                    return ('game_assets', 'fonts')

            # Check regex patterns for numbered sprites and variants
            for pattern in self.game_sprite_patterns:
                if pattern.match(clean_stem):
                    return ('game_assets', 'sprites')

            # Check for sprite/texture keyword patterns
            for keyword in self.game_sprite_keywords:
                if keyword in stem or keyword in clean_stem:
                    # Distinguish between sprites and textures
                    sprite_keywords = [
                        'frame', 'sprite', 'leg', 'arm', 'head', 'torso', 'body',
                        'wing', 'hair', 'face', 'mouth', '_grey', '_gray',
                        'assassins', 'atonement', 'arrow_v', 'arrow_h', 'add',
                        '2h_', '1h_', 'dagger', 'sword', 'axe', 'hammer', 'mace',
                        # Character customization
                        'beard', 'bling', 'hiero', 'mustache', 'scar', 'tattoo',
                        'earring', 'necklace', 'bracelet', 'glasses', 'mask', 'hood'
                    ]
                    if any(kw in stem or kw in clean_stem for kw in sprite_keywords):
                        return ('game_assets', 'sprites')
                    else:
                        return ('game_assets', 'textures')

        # Check for font files
        if ext in ['.ttf', '.otf', '.woff', '.woff2', '.eot', '.fon', '.fnt']:
            if ext == '.ttf':
                return ('fonts', 'truetype')
            elif ext == '.otf':
                return ('fonts', 'opentype')
            elif ext in ['.woff', '.woff2', '.eot']:
                return ('fonts', 'web')
            else:
                return ('fonts', 'other')

        return None

    def classify_by_organization(self, text: str, filename: str) -> Optional[Tuple[str, str, str]]:
        """
        Classify file primarily by Organization entity detection.

        Looks for strong organization indicators like:
        - Company names in headers/footers
        - Official letterheads
        - Business correspondence
        - Invoices, contracts with company names

        Returns:
            Tuple of (category, subcategory, org_name) or None if no strong organization match
        """
        if not text or len(text) < 50:
            return None

        text_lower = text.lower()
        filename_lower = filename.lower()

        # Organization type indicators
        org_indicators = {
            'government': [
                'department of', 'internal revenue', 'irs', 'social security',
                'state of', 'county of', 'city of', 'municipality', 'federal',
                'government', 'agency', 'bureau', 'commission', 'dmv',
                'passport', 'immigration', 'customs', 'treasury'
            ],
            'healthcare': [
                'hospital', 'clinic', 'medical center', 'health system',
                'healthcare', 'physicians', 'doctor', 'patient', 'diagnosis',
                'prescription', 'pharmacy', 'insurance claim', 'medicare',
                'medicaid', 'hipaa', 'medical record', 'lab results'
            ],
            'financial': [
                'bank', 'credit union', 'investment', 'brokerage', 'mortgage',
                'loan', 'account statement', 'transaction', 'wire transfer',
                'routing number', 'account number', 'fdic', 'securities'
            ],
            'educational': [
                'university', 'college', 'school', 'academy', 'institute',
                'transcript', 'diploma', 'degree', 'enrollment', 'registrar',
                'financial aid', 'tuition', 'semester', 'course', 'student id'
            ],
            'nonprofit': [
                'foundation', 'charity', 'nonprofit', 'non-profit', '501(c)',
                'donation', 'volunteer', 'mission', 'charitable'
            ],
            'employers': [
                'offer letter', 'employment agreement', 'w-2', 'w2', 'pay stub',
                'payroll', 'human resources', 'hr department', 'employee id',
                'benefits enrollment', 'performance review', 'termination'
            ],
            'vendors': [
                'invoice', 'purchase order', 'po number', 'vendor id',
                'supplier', 'bill to', 'ship to', 'payment terms', 'net 30'
            ],
            'clients': [
                'client', 'customer', 'service agreement', 'statement of work',
                'sow', 'proposal', 'quote', 'estimate', 'engagement letter'
            ]
        }

        # Check for organization type indicators
        for org_type, keywords in org_indicators.items():
            matches = sum(1 for kw in keywords if kw in text_lower)
            if matches >= 2:  # Require at least 2 keyword matches
                # Try to extract organization name
                companies = self.classifier.extract_company_names(text)
                org_name = companies[0] if companies else None
                if org_name:
                    return ('organization', org_type, org_name)

        return None

    def classify_by_person(self, text: str, filename: str) -> Optional[Tuple[str, str, List[str]]]:
        """
        Classify file primarily by Person entity detection.

        Looks for strong person indicators like:
        - Resumes/CVs
        - Contact information (vCards)
        - Personal identification documents
        - Reference letters

        Returns:
            Tuple of (category, subcategory, person_names) or None if no strong person match
        """
        if not text or len(text) < 50:
            return None

        text_lower = text.lower()
        filename_lower = filename.lower()

        # Person type indicators
        person_indicators = {
            'contacts': [
                'contact', 'phone:', 'email:', 'address:', 'mobile:',
                'tel:', 'fax:', 'linkedin', 'twitter', '@'
            ],
            'employees': [
                'employee', 'staff', 'team member', 'department:', 'title:',
                'hire date', 'start date', 'position:', 'role:'
            ],
            'references': [
                'reference', 'recommendation', 'letter of', 'to whom it may concern',
                'i am pleased to', 'i highly recommend', 'worked with'
            ],
            'clients': [
                'client profile', 'customer profile', 'client information',
                'account holder', 'policyholder'
            ]
        }

        # Check filename patterns for resumes/CVs
        resume_patterns = ['resume', 'cv', 'curriculum', 'vitae']
        if any(pat in filename_lower for pat in resume_patterns):
            people = self.classifier.extract_people_names(text)
            return ('person', 'contacts', people if people else [])

        # Check for person type indicators
        for person_type, keywords in person_indicators.items():
            matches = sum(1 for kw in keywords if kw in text_lower)
            if matches >= 2:  # Require at least 2 keyword matches
                people = self.classifier.extract_people_names(text)
                if people:
                    return ('person', person_type, people)

        return None

    def classify_media_file(self, file_path: Path, image_metadata: Dict = None) -> Optional[Tuple[str, str, str]]:
        """
        Classify media files (photos, videos, audio) into subcategories.

        Returns:
            Tuple of (category, media_type, subcategory) or None if not a media file
            Example: ('media', 'photos', 'screenshots') or ('media', 'videos', 'recordings')
        """
        filename = file_path.name.lower()
        stem = file_path.stem.lower()
        ext = file_path.suffix.lower()

        # Videos - .mp4, .mov, .avi, .mkv, .webm, .m4v
        if ext in ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.m4v', '.flv', '.wmv']:
            # Screen recordings
            if 'screen' in stem or 'recording' in stem or 'capture' in stem:
                return ('media', 'videos', 'screencasts')
            # Exports (from video editors)
            elif 'export' in stem or 'render' in stem or 'final' in stem or 'cut' in stem:
                return ('media', 'videos', 'exports')
            # Default to recordings
            else:
                return ('media', 'videos', 'recordings')

        # Audio - .mp3, .wav, .m4a, .aac, .flac, .ogg (but not game music)
        if ext in ['.mp3', '.m4a', '.aac', '.flac', '.wma']:
            # Podcasts
            if 'podcast' in stem or 'episode' in stem or 'interview' in stem:
                return ('media', 'audio', 'podcasts')
            # Music
            elif 'song' in stem or 'album' in stem or 'track' in stem or 'music' in stem:
                return ('media', 'audio', 'music')
            # Voice recordings
            elif 'recording' in stem or 'voice' in stem or 'memo' in stem or 'audio' in stem:
                return ('media', 'audio', 'recordings')
            # Default to recordings
            else:
                return ('media', 'audio', 'recordings')

        # Photos - .jpg, .jpeg, .png, .heic, .gif, .webp, .bmp
        if ext in ['.jpg', '.jpeg', '.png', '.heic', '.gif', '.webp', '.bmp', '.tiff', '.tif']:
            # Screenshots (highest priority for photos)
            if filename.startswith('screenshot') or 'screen shot' in filename:
                return ('media', 'photos', 'screenshots')

            # Scanned documents/receipts (OCR will detect text)
            if 'scan' in stem or 'receipt' in stem or 'document' in stem or 'invoice' in stem:
                return ('media', 'photos', 'documents')

            # Travel photos (has GPS metadata)
            if image_metadata and image_metadata.get('gps_coordinates'):
                # If we have GPS coordinates, it's likely a travel photo
                return ('media', 'photos', 'travel')

            # Photos with datetime (camera photos) - organize by type
            if image_metadata and image_metadata.get('datetime'):
                # Photos with camera EXIF data are likely personal photos
                # Default to 'other' category for general photos
                return ('media', 'photos', 'other')

            # Photos without metadata - still categorize as media if they're actual photos
            # (as opposed to game sprites which would be caught earlier)
            if ext in ['.jpg', '.jpeg', '.heic']:
                return ('media', 'photos', 'other')

            # PNG files without clear classification fall through
            # (could be screenshots, documents, or game assets that weren't caught)
            return None

        return None

    def classify_by_filename_patterns(self, file_path: Path) -> Optional[Tuple[str, str, Optional[str], List[str]]]:
        """
        Classify file based on filename patterns before content extraction.

        This method handles common patterns that can be quickly identified
        by filename alone, avoiding expensive OCR/content analysis.

        Returns:
            Tuple of (category, subcategory, company_name, people_names) or None
        """
        filename = file_path.name
        filename_lower = filename.lower()
        stem = file_path.stem.lower()
        ext = file_path.suffix.lower()
        original_stem = file_path.stem

        # =========================================================
        # DUPLICATE DETECTION: Files with _YYYYMMDD_HHMMSS suffix
        # These are timestamped copies - skip them
        # =========================================================
        if re.search(r'_\d{8}_\d{6}$', file_path.stem):
            print(f"  ⚠ Duplicate file (timestamped copy) - skipping")
            return ('skip', 'duplicate', None, [])

        # =========================================================
        # LOG FILES: System logs, reorganization logs → Technical/Logs
        # =========================================================
        log_patterns = ['reorganization-log', 'reorganization_log', 'system-log', 'system_log',
                       'error-log', 'error_log', 'debug-log', 'debug_log', 'access-log', 'access_log']
        if any(p in stem for p in log_patterns) or (ext == '.log'):
            print(f"  ✓ Filename pattern: Log file")
            return ('technical', 'logs', None, [])

        # =========================================================
        # LEGAL DOCUMENT PATTERNS: Contracts, Corporate docs
        # =========================================================
        # Certificate of Formation/Filing → Legal/Corporate
        corporate_legal_patterns = ['certificateofformation', 'certificate_of_formation', 'certificate-of-formation',
                                   'certificateoffiling', 'certificate_of_filing', 'certificate-of-filing',
                                   'articlesofincorporation', 'articles_of_incorporation', 'articles-of-incorporation',
                                   'bylaws', 'operatingagreement', 'operating_agreement', 'operating-agreement']
        if any(p in stem for p in corporate_legal_patterns):
            print(f"  ✓ Filename pattern: Corporate legal document")
            return ('legal', 'corporate', None, [])

        # Release of Liability, NDA, General Release → Legal/Contracts
        contract_legal_patterns = ['releaseofliability', 'release_of_liability', 'release-of-liability',
                                  'generalrelease', 'general_release', 'general-release',
                                  'non-disclosure', 'nondisclosure', 'confidentiality']
        # NDA needs special handling - must be whole word or at word boundary (not inside 'calendar')
        is_nda = (stem == 'nda' or stem.startswith('nda_') or stem.startswith('nda-') or
                  '_nda' in stem or '-nda' in stem or stem.endswith('_nda') or stem.endswith('-nda'))
        if any(p in stem for p in contract_legal_patterns) or is_nda:
            print(f"  ✓ Filename pattern: Legal contract/release")
            return ('legal', 'contracts', None, [])

        # =========================================================
        # TECHNICAL CONFIG: Login credentials, account recovery docs
        # =========================================================
        config_patterns = ['login', 'recovery', 'credentials', 'password', 'apikey', 'api_key', 'api-key',
                          'techsoup', 'zoho', 'oauth', 'token', 'secret']
        if ext in {'.docx', '.doc', '.txt', '.pdf'} and any(p in stem for p in config_patterns):
            print(f"  ✓ Filename pattern: Config/credentials document")
            return ('technical', 'config', None, [])

        # =========================================================
        # MARKETING PATTERNS: Strategy docs, market maps, infographics
        # =========================================================
        marketing_patterns = ['marketingstrategy', 'marketing_strategy', 'marketing-strategy',
                             'marketmap', 'market_map', 'market-map', 'marketanalysis', 'market_analysis',
                             'competitoranalysis', 'competitor_analysis', 'brandstrategy', 'brand_strategy',
                             'contentcalendar', 'content_calendar', 'socialmedia', 'social_media',
                             'infographic', 'info_graphic', 'info-graphic']
        if any(p in stem for p in marketing_patterns):
            print(f"  ✓ Filename pattern: Marketing document")
            return ('business', 'marketing', None, [])

        # =========================================================
        # COMPANY MEETING NOTES: Must be before generic company patterns
        # =========================================================
        # Integrity Studio meeting notes (IntegrityWeeklyCadence-*-NotesByGemini.docx)
        if 'integrityweeklycadence' in stem or ('integrity' in stem and 'cadence' in stem):
            print(f"  ✓ Filename pattern: Integrity Studio meeting notes")
            return ('organization', 'meeting_notes', 'Integrity Studio', [])

        # =========================================================
        # LEORA HOME HEALTH: Stock photos and business assets
        # =========================================================
        if ext in {'.jpg', '.jpeg', '.png', '.webp'}:
            leora_keywords = ['elderlycare', 'caregiver', 'home-health', 'homehealth',
                             'skilled-nursing', 'skillednursing', 'at-home-nurse',
                             'compassionate-home', 'grandparents-and', 'get-started-seniors',
                             'daughter-and-mother']
            leora_prefixes = ['atx-caregiver', 'atx-nurse']
            if any(kw in stem for kw in leora_keywords) or any(stem.startswith(p) for p in leora_prefixes):
                print(f"  ✓ Filename pattern: Leora Home Health asset")
                return ('organization', 'other', 'Leora Home Health', [])

        # =========================================================
        # COMPANY-BASED ORGANIZATION: Files with company names
        # =========================================================
        company_patterns = {
            'integrity': ('organization', 'other', 'Integrity Studio'),
            'integrityai': ('organization', 'other', 'Integrity Studio'),
            'integritystudio': ('organization', 'other', 'Integrity Studio'),
            'integrity_studio': ('organization', 'other', 'Integrity Studio'),
            'integrity-studio': ('organization', 'other', 'Integrity Studio'),
            'integritycrm': ('organization', 'other', 'Integrity Studio'),
            'inspiredmovement': ('organization', 'vendors', 'Inspired Movement'),
            'inspired_movement': ('organization', 'vendors', 'Inspired Movement'),
            'inspired-movement': ('organization', 'vendors', 'Inspired Movement'),
            'inspired': ('organization', 'vendors', 'Inspired Movement'),
        }
        for pattern, (category, subcategory, company_name) in company_patterns.items():
            if pattern in stem:
                print(f"  ✓ Filename pattern: Company file ({company_name})")
                return (category, subcategory, company_name, [])

        # =========================================================
        # BUSINESS TYPE PATTERNS: CRM, HR, Operations, Planning
        # Skip code files - they go to Technical even if they have business keywords
        # =========================================================
        # Code file extensions that should NOT be matched by business patterns
        # Includes web build artifacts (css, scss) that often have hashed names like client.063172a3.css
        business_skip_extensions = {'.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.go', '.rs', '.rb', '.php',
                                   '.css', '.scss', '.less', '.sass', '.html', '.htm', '.vue', '.svelte'}

        # CRM files (but not code files like contacts.py, essentialcontacts_v1_client.py)
        if ext not in business_skip_extensions:
            if 'crm' in stem or 'microlender' in stem:
                print(f"  ✓ Filename pattern: CRM/Contacts file")
                return ('business', 'crm', None, [])
            # Only match 'contacts' for spreadsheet/document files, not code
            if 'contacts' in stem and ext in {'.xlsx', '.xls', '.csv', '.docx', '.pdf'}:
                print(f"  ✓ Filename pattern: CRM/Contacts file")
                return ('business', 'crm', None, [])
            # Third party vendor/partner evaluation files
            thirdparty_patterns = ['3rdparties', '3rdparty', 'thirdparties', 'thirdparty',
                                   'third_parties', 'third_party', 'third-parties', 'third-party']
            if any(p in stem for p in thirdparty_patterns):
                print(f"  ✓ Filename pattern: Third-party vendor/partner list")
                return ('business', 'crm', None, [])

        # HR/Job posting files (but not code files like application.py, linkedin.py)
        # Note: 'application' and 'linkedin' are common in code filenames
        hr_patterns = ['jobposting', 'job_posting', 'job-posting',
                       'boardmember', 'board_member', 'hiring', 'teamroster', 'team_roster', 'team-roster']
        # Only match 'application' or 'linkedin' for document files
        hr_doc_patterns = ['application', 'linkedin']
        if ext not in business_skip_extensions:
            if any(p in stem for p in hr_patterns):
                print(f"  ✓ Filename pattern: HR file")
                return ('business', 'hr', None, [])
            if ext in {'.xlsx', '.xls', '.csv', '.docx', '.pdf', '.webp', '.png', '.jpg'} and any(p in stem for p in hr_doc_patterns):
                print(f"  ✓ Filename pattern: HR file")
                return ('business', 'hr', None, [])

        # Project tracking and product planning files (but not code files)
        if ext not in business_skip_extensions:
            if 'projecttrack' in stem or 'project_track' in stem or 'project-track' in stem:
                print(f"  ✓ Filename pattern: Project tracking")
                return ('business', 'planning', None, [])
            # Product ideas, analysis, and roadmap files
            planning_patterns = ['productideas', 'product_ideas', 'product-ideas',
                                'productanalysis', 'product_analysis', 'product-analysis',
                                'productroadmap', 'product_roadmap', 'product-roadmap',
                                'roadmap_product', 'roadmap-product', 'roadmapproduct']
            if any(p in stem for p in planning_patterns):
                print(f"  ✓ Filename pattern: Product planning/analysis")
                return ('business', 'planning', None, [])

        # Operations/Dashboard files (but not code files like dashboard.py, operations.py)
        if ext not in business_skip_extensions:
            if 'dashboard' in stem or 'operations' in stem or 'toolkit' in stem:
                print(f"  ✓ Filename pattern: Operations file")
                return ('business', 'other', None, [])

        # Stand-up/meeting templates and notes
        meeting_patterns = ['standup', 'stand-up', 'stand_up', 'meeting', 'minutes', 'agenda',
                           'allhands', 'all-hands', 'all_hands', 'retrospective', 'retro']
        if ext not in business_skip_extensions and any(p in stem for p in meeting_patterns):
            print(f"  ✓ Filename pattern: Meeting notes/template")
            return ('business', 'meeting_notes', None, [])

        # Shipping labels
        if 'printlabel' in stem or 'print_label' in stem or 'shippinglabel' in stem:
            print(f"  ✓ Filename pattern: Shipping label")
            return ('business', 'other', None, [])

        # =========================================================
        # GOOGLE INVOICES: 51xxxxx.pdf, 52xxxxx.pdf, 53xxxxx.pdf patterns
        # =========================================================
        if ext == '.pdf' and re.match(r'^5[123]\d{8,}', filename):
            print(f"  ✓ Filename pattern: Google invoice")
            return ('organization', 'vendors', 'Google', [])

        # =========================================================
        # SCREENSHOTS: Files with 'screenshot' in name → Media/Photos/Screenshots
        # =========================================================
        if 'screenshot' in stem:
            print(f"  ✓ Filename pattern: Screenshot")
            return ('media', 'photos_screenshots_other', None, [])

        # =========================================================
        # SURVEYS: Survey/questionnaire documents → Business/Other
        # =========================================================
        survey_patterns = ['survey', 'questionnaire', 'feedback-form', 'feedback_form']
        if ext in {'.docx', '.doc', '.pdf', '.xlsx', '.xls'} and any(p in stem for p in survey_patterns):
            print(f"  ✓ Filename pattern: Survey/questionnaire")
            return ('business', 'other', None, [])

        # Known person name patterns (used in multiple places)
        known_person_patterns = {
            'ledlie': 'Alyshia Ledlie',
            'alyshia': 'Alyshia Ledlie',
        }

        # =========================================================
        # RESUME FILES: *resume*, *cv* - ONLY for document file types
        # Code files (.py, .js, .ts, .css) with "resume" are technical files
        # =========================================================
        document_extensions = {'.pdf', '.docx', '.doc', '.rtf', '.odt', '.txt'}
        code_extensions = {'.py', '.js', '.ts', '.jsx', '.tsx', '.css', '.scss', '.less', '.html', '.htm'}

        resume_patterns = ['resume', 'curriculum_vitae', 'curriculum-vitae']
        # Match 'cv' if it's the whole filename, has separators, or starts with 'cv ' (space)
        stem_lower = stem.lower()
        is_cv_document = (stem_lower == 'cv' or stem_lower.startswith('cv_') or stem_lower.startswith('cv-') or
                         stem_lower.startswith('cv ') or '_cv' in stem_lower or '-cv' in stem_lower or
                         ' cv' in stem_lower or stem_lower.endswith('_cv') or stem_lower.endswith('-cv'))

        if ext in document_extensions and (any(p in filename_lower for p in resume_patterns) or is_cv_document):
            # Try to extract person name from filename
            person_name = None

            # Pattern 1: CV/Resume at START followed by FirstName LastName (e.g., "CV Isabel Budenz January")
            name_match = re.search(r'^(cv|resume)[_\-\s]+([A-Z][a-z]+)[_\-\s]+([A-Z][a-z]+)', filename, re.IGNORECASE)
            if name_match:
                person_name = f"{name_match.group(2)} {name_match.group(3)}"

            # Pattern 2: FirstName LastName at START followed by CV/Resume (e.g., "Alyshia_Ledlie_Technical_Resume")
            # Take the FIRST two capitalized words as the name
            if not person_name:
                name_match = re.search(r'^([A-Z][a-z]+)[_\-\s]+([A-Z][a-z]+)', filename)
                if name_match:
                    # Verify this file contains resume/cv somewhere
                    if 'resume' in filename_lower or is_cv_document:
                        candidate = f"{name_match.group(1)} {name_match.group(2)}"
                        # Skip template/style names that aren't person names
                        template_words = {'modern', 'minimalist', 'professional', 'creative', 'simple',
                                         'elegant', 'classic', 'template', 'standard', 'basic', 'clean'}
                        first_word = name_match.group(1).lower()
                        second_word = name_match.group(2).lower()
                        if first_word not in template_words and second_word not in template_words:
                            person_name = candidate

            # Check for known person names if no name extracted yet
            if not person_name:
                for pattern, known_name in known_person_patterns.items():
                    if pattern in filename_lower:
                        person_name = known_name
                        break

            # If name found in filename, return immediately
            if person_name:
                print(f"  ✓ Filename pattern: Resume ({person_name})")
                return ('person', 'contacts', None, [person_name])
            # If no name found, don't return - fall through to OCR content extraction
            # This allows "Modern Minimalist CV Resume.pdf" to get name from PDF content
            print(f"  ✓ Filename pattern: Resume (extracting name from content...)")

        # =========================================================
        # COVER LETTERS: *coverletter*, *cover_letter* - document files
        # =========================================================
        cover_letter_patterns = ['coverletter', 'cover_letter', 'cover-letter']
        if ext in document_extensions and any(p in stem for p in cover_letter_patterns):
            # Try to extract person name from filename
            person_name = None
            for pattern, known_name in known_person_patterns.items():
                if pattern in filename_lower:
                    person_name = known_name
                    break
            print(f"  ✓ Filename pattern: Cover letter" + (f" ({person_name})" if person_name else ""))
            return ('person', 'contacts', None, [person_name] if person_name else [])

        # Check for known person names in filename (e.g., ledlie) - non-resume files
        for pattern, person_name in known_person_patterns.items():
            if pattern in filename_lower:
                print(f"  ✓ Filename pattern: Person ({person_name})")
                return ('person', 'contacts', None, [person_name])

        # =========================================================
        # ENTITY-BASED FILES: Company names in filename (check BEFORE extension)
        # This ensures files like LeoraHomeHealth-Data.csv go to Organization/
        # rather than Technical/
        # =========================================================
        entity_patterns = {
            'integritystudio': ('organization', 'vendors', 'Integrity Studio'),
            'integrity_studio': ('organization', 'vendors', 'Integrity Studio'),
            'integrity-studio': ('organization', 'vendors', 'Integrity Studio'),
            'leora': ('organization', 'healthcare', 'Leora Home Health'),
            'leorahomehealth': ('organization', 'healthcare', 'Leora Home Health'),
            'ltchcssa': ('organization', 'healthcare', 'Leora Home Health'),  # LTCHCSSA is Leora related
            'inspiredmovement': ('organization', 'vendors', 'Inspired Movement'),
            'inspired_movement': ('organization', 'vendors', 'Inspired Movement'),
            'inspired-movement': ('organization', 'vendors', 'Inspired Movement'),
            'fisterra': ('organization', 'vendors', 'Fisterra'),
            'dotfun': ('organization', 'vendors', 'DotFun'),
            'ensco': ('organization', 'vendors', 'EnsoCo'),
            'capitalcityvillage': ('organization', 'property_management', 'Capital City Village'),
            'capital_city_village': ('organization', 'property_management', 'Capital City Village'),
            'capital-city-village': ('organization', 'property_management', 'Capital City Village'),
            'google': ('organization', 'vendors', 'Google'),
            'microsoft': ('organization', 'vendors', 'Microsoft'),
            'adobe': ('organization', 'vendors', 'Adobe Systems'),
            'amazon': ('organization', 'vendors', 'Amazon'),
            'apple': ('organization', 'vendors', 'Apple'),
            'genius bar': ('organization', 'vendors', 'Apple'),
        }
        for pattern, (category, subcat, company_name) in entity_patterns.items():
            if pattern in stem:
                print(f"  ✓ Filename pattern: Entity ({company_name})")
                return (category, subcat, company_name, [])

        # =========================================================
        # TECHNICAL FILES: .py, .js, .ts, .csv, .json, .xml
        # (Only if no entity pattern matched above)
        # =========================================================
        technical_extensions = {
            '.py': 'Technical',
            '.js': 'Technical',
            '.ts': 'Technical',
            '.jsx': 'Technical',
            '.tsx': 'Technical',
            '.csv': 'Technical',
            '.json': 'Technical',
            '.xml': 'Technical',
            '.yaml': 'Technical',
            '.yml': 'Technical',
            '.sql': 'Technical',
            '.sh': 'Technical',
            '.bash': 'Technical',
        }
        if ext in technical_extensions:
            print(f"  ✓ Filename pattern: Technical file ({ext})")
            return ('technical', 'other', None, [])

        # =========================================================
        # SOFTWARE PACKAGES: .dmg, .pkg, .msi, .deb, .rpm, .exe, .app
        # =========================================================
        software_extensions = {'.dmg', '.pkg', '.msi', '.deb', '.rpm', '.exe',
                               '.app', '.snap', '.flatpak', '.appimage'}
        if ext in software_extensions:
            print(f"  ✓ Filename pattern: Software package ({ext})")
            return ('technical', 'software_packages', None, [])

        # =========================================================
        # LEGAL DOCUMENTS: Agreement, CLA, Operating, Reseller, Severance
        # =========================================================
        legal_patterns = [
            ('agreement', 'contracts'),
            ('operating', 'corporate'),
            ('reseller', 'contracts'),
            ('severance', 'contracts'),
            ('contract', 'contracts'),
            ('amendment', 'contracts'),
            ('certificateoffiling', 'corporate'),
        ]
        for pattern, subcat in legal_patterns:
            if pattern in stem:
                print(f"  ✓ Filename pattern: Legal document ({pattern})")
                return ('legal', subcat, None, [])

        # CLA pattern - Contributor License Agreement
        # Be careful not to match "class", "clause", "claw", "eucla" (timezone), etc.
        # Only match explicit CLA patterns
        cla_patterns = [
            r'^cla[_\-\d.]',              # cla_signed, cla-2024, cla.pdf
            r'^cla$',                      # just "cla"
            r'corporatecla',               # CorporateCLA
            r'individualcla',              # IndividualCLA
            r'contributorcla',             # ContributorCLA
            r'[_\-]cla[_\-\d.]',          # some_cla_file, my-cla-2024
            r'[_\-]cla$',                  # some_cla, my-cla
        ]
        if any(re.search(p, stem) for p in cla_patterns):
            print(f"  ✓ Filename pattern: Legal document (CLA)")
            return ('legal', 'contracts', None, [])

        # =========================================================
        # BUSINESS DOCUMENTS: BizAid, BizStart, Meeting
        # =========================================================
        business_patterns = [
            ('bizaid', 'planning'),
            ('bizstart', 'planning'),
            ('meeting', 'other'),
            ('proposal', 'proposals'),
        ]
        for pattern, subcat in business_patterns:
            if pattern in stem:
                print(f"  ✓ Filename pattern: Business document ({pattern})")
                return ('business', subcat, None, [])

        # =========================================================
        # DATA USAGE AGREEMENTS: Special case
        # =========================================================
        if 'datausageagreement' in stem or 'data_usage_agreement' in stem or 'data-usage-agreement' in stem:
            print(f"  ✓ Filename pattern: Data Usage Agreement")
            return ('legal', 'contracts', None, [])

        # =========================================================
        # DOCUMENTATION FILES: LICENSE, README, specs
        # =========================================================
        # LICENSE files
        if stem.startswith('license') or stem == 'copying' or stem == 'licence':
            print(f"  ✓ Filename pattern: License file")
            return ('technical', 'documentation', None, [])

        # README files
        if stem.startswith('readme') or stem == 'read_me' or stem == 'read-me':
            print(f"  ✓ Filename pattern: README file")
            return ('technical', 'documentation', None, [])

        # Specification/spec documents
        spec_patterns = ['specification', 'spec_', '_spec', '-spec', 'specs_', '_specs']
        if any(p in stem for p in spec_patterns) or stem == 'spec' or stem == 'specs':
            print(f"  ✓ Filename pattern: Specification document")
            return ('technical', 'documentation', None, [])

        # CHANGELOG files
        if stem.startswith('changelog') or stem == 'changes' or stem == 'history':
            print(f"  ✓ Filename pattern: Changelog file")
            return ('technical', 'documentation', None, [])

        # =========================================================
        # COVER LETTERS: Go to Person folder
        # =========================================================
        if 'coverletter' in stem or 'cover_letter' in stem or 'cover-letter' in stem:
            # Try to extract person name
            person_name = None
            for pattern, known_name in known_person_patterns.items():
                if pattern in filename_lower:
                    person_name = known_name
                    break
            print(f"  ✓ Filename pattern: Cover letter" + (f" ({person_name})" if person_name else ""))
            return ('person', 'contacts', None, [person_name] if person_name else [])

        # =========================================================
        # CONFIG/MANIFEST FILES: Technical config
        # =========================================================
        config_patterns = ['.manifest', '.config', '.ini', '.cfg', '.conf', '.plist']
        if ext in config_patterns:
            print(f"  ✓ Filename pattern: Config file ({ext})")
            return ('technical', 'config', None, [])

        # Config-like text files (settings.txt, config.txt, preferences.txt)
        config_txt_names = ['settings', 'config', 'preferences', 'options', 'configuration']
        if ext == '.txt' and stem in config_txt_names:
            print(f"  ✓ Filename pattern: Config text file ({stem}.txt)")
            return ('technical', 'config', None, [])

        # =========================================================
        # GIT HOOK SAMPLES: .sample files from .git/hooks/
        # =========================================================
        if ext == '.sample':
            print(f"  ✓ Filename pattern: Git hook sample")
            return ('technical', 'config', None, [])

        # =========================================================
        # DOTFILES: .eslintrc, .editorconfig, .nycrc, .travis.yml, etc.
        # =========================================================
        if filename.startswith('.'):
            dotfile_configs = ['.eslintrc', '.editorconfig', '.nycrc', '.travis', '.codecov',
                              '.codeclimate', '.yarnrc', '.npmrc', '.prettierrc', '.babelrc',
                              '.gitignore', '.gitattributes', '.dockerignore', '.env']
            if any(filename.startswith(cfg) for cfg in dotfile_configs):
                print(f"  ✓ Filename pattern: Dotfile config ({filename})")
                return ('technical', 'config', None, [])
            # Generic dotfiles with common config extensions
            if ext in ['.yml', '.yaml', '.json', '.toml']:
                print(f"  ✓ Filename pattern: Dotfile config ({filename})")
                return ('technical', 'config', None, [])

        # =========================================================
        # SOURCE MAP FILES: .map, .js.map, .d.ts.map → Technical/Config
        # =========================================================
        if ext == '.map' or filename.endswith('.js.map') or filename.endswith('.d.ts.map'):
            print(f"  ✓ Filename pattern: Source map file")
            return ('technical', 'config', None, [])

        # =========================================================
        # MEETING NOTES: Generic meeting notes → Business/Other
        # (Company-specific meeting notes handled earlier in COMPANY MEETING NOTES section)
        # =========================================================
        meeting_patterns = ['weeklynotes', 'weekly_notes', 'weekly-notes',
                           'meetingnotes', 'meeting_notes', 'meeting-notes', 'notesby']
        if any(p in stem for p in meeting_patterns):
            print(f"  ✓ Filename pattern: Meeting notes")
            return ('business', 'other', None, [])

        # =========================================================
        # DOCUMENTATION/NOTES FILES: Websites, personal notes
        # =========================================================
        if stem.startswith('websites') or 'ivemade' in stem or "i'vemade" in stem.replace("'", "'"):
            print(f"  ✓ Filename pattern: Personal documentation")
            return ('business', 'other', None, [])

        # =========================================================
        # AUDIO FILES: Check for game audio first, then Media/Audio
        # =========================================================
        audio_extensions = {'.wav', '.ogg', '.mp3', '.flac', '.aac', '.m4a', '.wma'}
        if ext in audio_extensions:
            # Game audio keywords - sound effects and music
            # Note: Avoid 'cast' as it matches 'podcast'
            game_audio_keywords = [
                'bolt', 'spell', 'magic', 'spellcast', 'chirp', 'crossbow', 'dagger',
                'sword', 'arrow', 'bow', 'heal', 'potion', 'lightning', 'fire',
                'ice', 'acid', 'poison', 'explosion', 'blast', 'summon', 'dispel',
                'petrification', 'neutralize', 'slow', 'darkness', 'achievement',
                'quest', 'unlock', 'lock', 'door', 'chest', 'coin', 'pickup',
                'attack', 'hit', 'damage', 'death', 'footstep', 'jump', 'land',
                'monster', 'creature', 'enemy', 'boss', 'battle', 'combat',
                'starving', 'hunger', 'thirst', 'eat', 'drink', 'sleep',
                'fiddle', 'lute', 'mandoline', 'glockenspiel', 'instrument',
                'identify', 'greater', 'mental', 'melee', 'axe', 'mace', 'whip',
                'sabre', 'staff', 'thunder', 'confusion', 'telekinetic', 'mind',
                'cure', 'light', 'firebolt', 'fireball', 'boomer', 'skur',
                # Game music keywords
                'dungeon', 'castle', 'forest', 'town', 'village', 'temple',
                'ruins', 'cave', 'mountain', 'ocean', 'desert', 'snow', 'victory',
                'defeat', 'theme', 'menu', 'credits', 'intro', 'outro', 'mysterious',
                'dark', 'epic', 'calm', 'peaceful', 'tension', 'chaos', 'hope',
                'despair', 'triumph', 'march', 'symphony', 'monotony', 'drakalor',
                'altar', 'lawful', 'chaotic', 'dwarven', 'elven', 'orcish', 'halls',
                'abandon', 'corrupting', 'breeze', 'clockwork', 'knowledge', 'final',
                'welcome', 'khelavaster', 'prophecy', 'spiraling', 'stairs', 'windy',
                'growth', 'warm', 'folk', 'peace', 'heritage', 'browsing', 'dusty',
                'tomes', 'dead', 'silent', 'ancardia', 'goblin', 'drums', 'bird'
            ]
            stem_lower = stem.lower()
            for keyword in game_audio_keywords:
                if keyword in stem_lower:
                    print(f"  ✓ Filename pattern: Game audio file ({stem}{ext}) → GameAssets/Audio")
                    return ('game_assets', 'audio', None, [])
            # Default to Media/Audio for non-game audio
            print(f"  ✓ Filename pattern: Audio file ({ext})")
            return ('media', 'audio_other', None, [])

        # =========================================================
        # VECTOR GRAPHICS: .svg, .ai, .eps → Media/Graphics
        # =========================================================
        vector_extensions = {'.svg', '.ai', '.eps'}
        if ext in vector_extensions:
            print(f"  ✓ Filename pattern: Vector graphics ({ext})")
            return ('media', 'graphics_vector', None, [])

        # =========================================================
        # GAME DATA FILES: .noe (ADOM game data), etc → GameAssets
        # =========================================================
        game_data_extensions = {'.noe'}
        if ext in game_data_extensions:
            print(f"  ✓ Filename pattern: Game data file ({ext})")
            return ('game_assets', 'other', None, [])

        # =========================================================
        # DIAGRAM/DOCUMENTATION IMAGES: *Diagram*, *ClassDiagram*
        # =========================================================
        if ext in {'.png', '.jpg', '.jpeg', '.webp', '.gif'}:
            if 'diagram' in stem or 'classdiagram' in stem:
                print(f"  ✓ Filename pattern: Diagram image")
                return ('technical', 'documentation', None, [])

        # =========================================================
        # GAME ASSET SPRITES: claw_*, icon_class_*, icon_* prefixes
        # =========================================================
        if ext in {'.png', '.jpg', '.jpeg', '.webp', '.gif'}:
            game_sprite_prefixes = ['claw_', 'icon_class_', 'icon_', 'sword_', 'shield_',
                                   'armor_', 'weapon_', 'item_', 'enemy_', 'player_',
                                   'tile_', 'bg_', 'effect_', 'spell_', 'skill_']
            if any(stem.startswith(p) for p in game_sprite_prefixes):
                print(f"  ✓ Filename pattern: Game sprite (prefix)")
                return ('game_assets', 'sprites', None, [])
            # Pattern: animation frames (2frame01.png, 3frame05.png)
            if re.match(r'^\d+frame\d+$', stem):
                print(f"  ✓ Filename pattern: Animation frame")
                return ('game_assets', 'sprites', None, [])
            # Pattern: BrogueFont files (BrogueFont1.png, etc.)
            if re.match(r'^broguefont\d+$', stem):
                print(f"  ✓ Filename pattern: Brogue font")
                return ('game_assets', 'fonts', None, [])

        # =========================================================
        # NUMBERED SPRITE FILES: 0.png, 103.png, 0_timestamp.png
        # Common pattern for game sprite sheets
        # =========================================================
        if ext in {'.png', '.jpg', '.jpeg', '.webp', '.gif', '.jp2'}:
            # EXCLUDE camera photo naming conventions from game asset detection
            # These are photos from cameras/phones, not game assets
            camera_prefixes = ('pxl_', 'img_', 'dsc_', 'dcim_', 'dscn_', 'dscf_', 'p_', 'photo_')
            is_camera_photo = stem.startswith(camera_prefixes)

            # EXCLUDE screenshot renamer software prefixes from game asset detection
            # These are software screenshots renamed by screenshot_renamer.py
            # Pattern: prefix_8hexchars (e.g., terminal_7a7c3ac6, browser_826a2e1f)
            software_screenshot_pattern = re.match(
                r'^(dashboard|terminal|code|browser|chat|settings|shop|product|docs|landing|infographic)_[a-f0-9]{8}$',
                stem
            )
            is_software_screenshot = software_screenshot_pattern is not None

            # Pattern: purely numeric filename (0.png, 103.png, 42_8.png)
            if not is_camera_photo and re.match(r'^\d+(_\d+)*$', stem):
                print(f"  ✓ Filename pattern: Numbered sprite")
                return ('game_assets', 'sprites', None, [])
            # Pattern: numeric with timestamp suffix (103_20251120_164958.png)
            if not is_camera_photo and re.match(r'^\d+(_\d+)*_\d{8}_\d{6}$', stem):
                print(f"  ✓ Filename pattern: Numbered sprite (timestamped)")
                return ('game_assets', 'sprites', None, [])
            # Pattern: font-related files - check BEFORE generic sprite patterns
            # (ascii_font.png, unicode_font.png, game_font.png, pixel_font.png)
            if not is_camera_photo and ('font' in stem or 'glyph' in stem or 'charset' in stem):
                print(f"  ✓ Filename pattern: Font asset")
                return ('game_assets', 'fonts', None, [])
            # Pattern: software screenshot from screenshot_renamer.py
            if is_software_screenshot:
                if stem.startswith('dashboard_'):
                    print(f"  ✓ Filename pattern: Software dashboard screenshot")
                    return ('media', 'photos_screenshots_dashboard', None, [])
                elif stem.startswith('terminal_'):
                    print(f"  ✓ Filename pattern: Terminal screenshot")
                    return ('media', 'photos_screenshots_terminal', None, [])
                elif stem.startswith('browser_'):
                    print(f"  ✓ Filename pattern: Browser screenshot")
                    return ('media', 'photos_screenshots_browser', None, [])
                elif stem.startswith('code_'):
                    print(f"  ✓ Filename pattern: Code editor screenshot")
                    return ('media', 'photos_screenshots_code', None, [])
                elif stem.startswith('docs_'):
                    print(f"  ✓ Filename pattern: Documentation screenshot")
                    return ('media', 'photos_screenshots_docs', None, [])
                elif stem.startswith(('shop_', 'product_')):
                    print(f"  ✓ Filename pattern: Product screenshot")
                    return ('media', 'photos_screenshots_products', None, [])
                elif stem.startswith('chat_'):
                    print(f"  ✓ Filename pattern: Chat screenshot")
                    return ('media', 'photos_screenshots_chat', None, [])
                elif stem.startswith('settings_'):
                    print(f"  ✓ Filename pattern: Settings screenshot")
                    return ('media', 'photos_screenshots_settings', None, [])
                elif stem.startswith(('landing_', 'infographic_')):
                    print(f"  ✓ Filename pattern: Marketing screenshot")
                    return ('business', 'marketing', None, [])
            # Pattern: name_hash.png (animal_57886bff.png, drop_2_6.png)
            # Exclude files where original stem has uppercase — those are human-named, not generated assets
            has_uppercase = any(c.isupper() for c in original_stem)
            if not is_camera_photo and not is_software_screenshot and not has_uppercase and re.match(r'^[a-z]+(_[a-z0-9]+)+$', stem):
                print(f"  ✓ Filename pattern: Game asset (named)")
                return ('game_assets', 'sprites', None, [])
            # Pattern: _hash or _name (starts with underscore, like _RWOIsUgWGL.png)
            if not is_camera_photo and re.match(r'^_[A-Za-z0-9]+(_\d{8}_\d{6})?$', stem):
                print(f"  ✓ Filename pattern: Game asset (underscore prefix)")
                return ('game_assets', 'sprites', None, [])
            # Pattern: number_word (17_in.png, 17_out_1.png)
            if not is_camera_photo and re.match(r'^\d+_[a-z]+(_\d+)?$', stem):
                print(f"  ✓ Filename pattern: Game asset (numbered)")
                return ('game_assets', 'sprites', None, [])
            # Pattern: hex codes for emoji/unicode (1f4a8.png, 1f600.png)
            if not is_camera_photo and re.match(r'^[0-9a-f]{4,8}$', stem):
                print(f"  ✓ Filename pattern: Emoji/unicode asset")
                return ('game_assets', 'sprites', None, [])
            # Pattern: date_category[_status]_id (ML training data)
            # Matches: 20190129_art_uncertain_100453.png, 20190129_pet_100453_1.png
            if not is_camera_photo and re.match(r'^\d{8}_[a-z]+(_[a-z]+)?_\d+(_\d+)*$', stem):
                print(f"  ✓ Filename pattern: ML training data")
                return ('game_assets', 'sprites', None, [])
            # Pattern: Facebook photo (481566579_10162021550590804_5823185318886800843_n.png)
            if re.match(r'^\d+_\d+_\d+_n$', stem):
                print(f"  ✓ Filename pattern: Social media photo")
                return ('media', 'photos_social', None, [])
            # Pattern: single word lowercase (achiever.png, sword.png)
            # Exclude data visualization and analytics terms
            data_viz_terms = {'pricing', 'trace', 'chart', 'graph', 'data', 'analytics',
                              'report', 'metrics', 'dashboard', 'distribution', 'histogram',
                              'timeline', 'funnel', 'heatmap', 'treemap', 'scatter', 'trend',
                              'forecast', 'summary', 'overview', 'statistics', 'benchmark',
                              'rework'}
            branding_terms = {'logo', 'logos', 'logotype', 'favicon', 'brandmark', 'wordmark'}
            portrait_terms = {'profile', 'headshot', 'portrait', 'avatar'}
            # Generic download names — no info about content, let fall through to CLIP/OCR
            generic_download_names = {'unnamed', 'untitled', 'image', 'photo', 'picture', 'screenshot', 'download'}
            # Tech product substrings — compound words describing real hardware, not game assets
            tech_product_substrings = {'macbook', 'iphone', 'ipad', 'airpods', 'android', 'laptop', 'tablet', 'keyboard', 'monitor', 'charger', 'adapter'}
            is_tech_product = any(t in stem for t in tech_product_substrings)
            if re.match(r'^[a-z]+$', stem) and len(stem) > 2 and stem not in data_viz_terms and stem not in branding_terms and stem not in portrait_terms and stem not in generic_download_names and not is_tech_product:
                print(f"  ✓ Filename pattern: Game asset (single word)")
                return ('game_assets', 'sprites', None, [])
            # Pattern: data visualization single word
            if re.match(r'^[a-z]+$', stem) and stem in data_viz_terms:
                print(f"  ✓ Filename pattern: Data visualization")
                return ('technical', 'data_visualization', None, [])
            # Pattern: ChatGPT images (ChatGPTImageNov1,2025,01_49_23AM.png)
            if stem.startswith('chatgptimage'):
                print(f"  ✓ Filename pattern: ChatGPT AI-generated image")
                return ('media', 'photos_chatgpt', None, [])
            # Pattern: Facebook images (481566579_10162021550590804_5823185318886800843_n.png)
            if re.match(r'^\d+_\d+_\d+_n$', stem):
                print(f"  ✓ Filename pattern: Facebook image")
                return ('media', 'photos_facebook', None, [])
            # Pattern: Portrait/profile photos (profile.png, headshot.png, avatar.png)
            if stem in portrait_terms or stem.startswith('profile') or stem.startswith('headshot'):
                print(f"  ✓ Filename pattern: Portrait photo")
                return ('media', 'photos_portraits', None, [])
            # Pattern: Logo images (logo-..., logotype-..., *-logo.png)
            if 'logo' in stem or 'logotype' in stem:
                print(f"  ✓ Filename pattern: Logo image")
                return ('organization', 'other', 'Integrity Studio', [])
            # Pattern: Leora Home Health stock photos (LHH-OG-*, nurse-*, medical-*)
            if stem.startswith('lhh-') or stem.startswith('lhh_'):
                print(f"  ✓ Filename pattern: Leora Home Health asset")
                return ('organization', 'healthcare', 'Leora Home Health', [])
            # Pattern: font-size files (courier-16.png, cp437-14_1.png, fantasy-16s.png)
            if re.match(r'^[a-z0-9]+-\d+[a-z]?(_\d+)?$', stem):
                print(f"  ✓ Filename pattern: Font/glyph file")
                return ('game_assets', 'fonts', None, [])
            # Pattern: codepage/charset font files (cp437-wide.png, cp437_wide_1.png, cp850-thin.png)
            if re.match(r'^cp\d+[_\-]', stem) or re.match(r'^(ascii|unicode|charset|codepage)[_\-]', stem):
                print(f"  ✓ Filename pattern: Codepage font file")
                return ('game_assets', 'fonts', None, [])
            # Pattern: game asset keywords with numbers (dungeon2, kitchen4, lightning1, interface2)
            # These are common game environment/effect/UI asset names
            game_asset_keywords = ['dungeon', 'kitchen', 'lightning', 'interface', 'items',
                                   'terinyo', 'castle', 'forest', 'cave', 'temple', 'tower',
                                   'weapon', 'armor', 'potion', 'scroll', 'effect', 'particle',
                                   'enemy', 'monster', 'creature', 'npc', 'player', 'character']
            for keyword in game_asset_keywords:
                if stem.startswith(keyword) and re.match(rf'^{keyword}\d+$', stem):
                    print(f"  ✓ Filename pattern: Game asset ({keyword})")
                    return ('game_assets', 'sprites', None, [])
            # Pattern: mixed case hash/ID (fSpW8I2Dxe6.png)
            if re.match(r'^[a-zA-Z0-9]{8,}$', stem) and not stem.isdigit() and not stem.isalpha():
                print(f"  ✓ Filename pattern: Hash/ID image")
                return ('media', 'photos_other', None, [])
            # Pattern: portrait/headshot photos (Name-p-800.jpg, Monica-p-1080.png)
            # Single capitalized word with -p-size suffix - typically named portrait photos
            if re.match(r'^[A-Z][a-z]+-p-\d+$', file_path.stem):  # Use original stem for case
                print(f"  ✓ Filename pattern: Portrait photo")
                return ('media', 'photos_portraits', None, [])
            # Pattern: hyphenated long names (stock photos with -p-500/-800/-1080 suffix)
            if re.match(r'^[a-z]+-[a-z]+-[a-z]+.*-p-\d+$', stem):
                print(f"  ✓ Filename pattern: Stock photo")
                return ('media', 'photos_stock', None, [])
            # Pattern: hyphenated long names without -p- suffix (general descriptive names)
            if re.match(r'^[a-z]+-[a-z]+-[a-z]+.*-\d+$', stem):
                print(f"  ✓ Filename pattern: Stock photo")
                return ('media', 'photos_stock', None, [])
            # Pattern: word_word (airbnb_earnings, austin_to_bombay)
            if re.match(r'^[a-z]+(_[a-z]+)+$', stem) and '_' in stem:
                print(f"  ✓ Filename pattern: Named image")
                return ('media', 'photos_other', None, [])
            # Pattern: letter+number sprites (l10.png, l20_1.png, note01.png, img1.png)
            # Exclude uppercase-origin filenames — human-named files, not generated game assets
            if not has_uppercase and re.match(r'^[a-z]+\d+(_\d+)?$', stem):
                print(f"  ✓ Filename pattern: Sprite sequence")
                return ('game_assets', 'sprites', None, [])
            # Pattern: word+number (drake2.png, grave2.png, void2.png)
            if not has_uppercase and re.match(r'^[a-z]+\d$', stem):
                print(f"  ✓ Filename pattern: Numbered variant")
                return ('game_assets', 'sprites', None, [])
            # Pattern: hyphenated names (heart-beat.png, phone-call.png)
            # Exclude data visualization hyphenated terms
            data_viz_hyphenated = {'yearly-distribution', 'monthly-distribution', 'daily-distribution',
                                   'cost-breakdown', 'revenue-chart', 'sales-report', 'time-series',
                                   'bar-chart', 'pie-chart', 'line-graph', 'data-flow', 'user-stats'}
            if re.match(r'^[a-z]+-[a-z]+(-[a-z]+)*(_\d{8}_\d{6})?(-copy)?$', stem):
                if stem in data_viz_hyphenated or any(term in stem for term in ['distribution', 'chart', 'graph', 'report', 'stats', 'analytics', 'metrics']):
                    print(f"  ✓ Filename pattern: Data visualization")
                    return ('technical', 'data_visualization', None, [])
                print(f"  ✓ Filename pattern: Hyphenated asset")
                return ('game_assets', 'sprites', None, [])
            # Pattern: two letters (dv.png, pv.png)
            if re.match(r'^[a-z]{2}$', stem):
                print(f"  ✓ Filename pattern: Two-letter asset")
                return ('game_assets', 'sprites', None, [])
            # Pattern: tinyfont/font with numbers (tinyfont66_1.png)
            if re.match(r'^[a-z]+font\d+(_\d+)?$', stem):
                print(f"  ✓ Filename pattern: Font sprite")
                return ('game_assets', 'fonts', None, [])
            # Pattern: repository templates
            if 'repository' in stem or 'template' in stem:
                print(f"  ✓ Filename pattern: Template image")
                return ('technical', 'other', None, [])
            # Pattern: requests-logo type (word-word.png)
            if re.match(r'^[a-z]+-[a-z]+(-compressed)?$', stem):
                print(f"  ✓ Filename pattern: Logo/brand image")
                return ('media', 'photos_other', None, [])
            # Pattern: hash ID with prefix (rs=xxx, shirt-xxx)
            if '=' in stem or re.match(r'^[a-z]+-\d+-[a-z0-9]+', stem):
                print(f"  ✓ Filename pattern: Generated ID image")
                return ('media', 'photos_other', None, [])

        # =========================================================
        # ICON FILES: .ico → Technical/Config, .icns → GameAssets/Other
        # =========================================================
        # favicon.ico and other .ico files are typically web config
        if ext == '.ico':
            print(f"  ✓ Filename pattern: Icon file ({stem}.ico) → Technical/Config")
            return ('technical', 'config', None, [])
        # .icns files are Mac app icons, often game-related
        if ext == '.icns':
            print(f"  ✓ Filename pattern: Mac icon file ({stem}.icns) → GameAssets/Other")
            return ('game_assets', 'other', None, [])

        # =========================================================
        # ARCHIVE FILES: .zip, .tar, .gz, .rar → Technical/Archives
        # =========================================================
        archive_extensions = {'.zip', '.tar', '.gz', '.rar', '.7z', '.bz2'}
        if ext in archive_extensions:
            print(f"  ✓ Filename pattern: Archive file ({ext})")
            return ('technical', 'archives', None, [])

        # =========================================================
        # CERTIFICATE/KEY FILES: .pem, .crt, .key → Technical/Security
        # =========================================================
        cert_extensions = {'.pem', '.crt', '.key', '.cer', '.p12', '.pfx'}
        if ext in cert_extensions:
            print(f"  ✓ Filename pattern: Certificate/key file ({ext})")
            return ('technical', 'security', None, [])

        # =========================================================
        # TEMPLATE FILES: .tpl → Technical/Templates
        # =========================================================
        if ext == '.tpl':
            print(f"  ✓ Filename pattern: Template file")
            return ('technical', 'templates', None, [])

        # =========================================================
        # FILES WITHOUT EXTENSION: Likely system/timezone data
        # =========================================================
        if not ext:
            # City/location names (timezone data): Abidjan, Accra, Adelaide, BajaNorte
            if re.match(r'^[A-Z][a-zA-Z-]+$', filename):
                print(f"  ✓ Filename pattern: Timezone/system data")
                return ('technical', 'other', None, [])
            # ALL CAPS names or with underscores (system data): ACT, ADOM, ADOM_1
            if re.match(r'^[A-Z][A-Z0-9_]+$', filename):
                print(f"  ✓ Filename pattern: System data")
                return ('technical', 'other', None, [])
            # Numeric IDs (social media): 2242610712719705
            if re.match(r'^\d{10,}$', filename):
                print(f"  ✓ Filename pattern: Numeric ID file")
                return ('technical', 'other', None, [])
            # Hash strings: 93419027627913a58f0b3fbb9ba9decea2a6bb
            if re.match(r'^[0-9a-f]{20,}$', filename):
                print(f"  ✓ Filename pattern: Hash file")
                return ('technical', 'other', None, [])
            # Script/tool names: activate-global-python-argcomplete
            if re.match(r'^[a-z]+(-[a-z]+)+$', filename):
                print(f"  ✓ Filename pattern: Script/tool")
                return ('technical', 'other', None, [])
            # CamelCase system files: CodeDirectory, CodeResources
            if re.match(r'^[A-Z][a-z]+[A-Z][a-zA-Z0-9-]*$', filename):
                print(f"  ✓ Filename pattern: macOS system file")
                return ('technical', 'other', None, [])
            # Lowercase with underscore and number: tsserver_1, bq
            if re.match(r'^[a-z]+(_\d+)?$', filename):
                print(f"  ✓ Filename pattern: System tool")
                return ('technical', 'other', None, [])
            # GMT timezones: GMT-0, GMT-1, GMT+5
            if re.match(r'^(GMT|UTC)[+-]?\d+$', filename):
                print(f"  ✓ Filename pattern: Timezone data")
                return ('technical', 'other', None, [])
            # ChangeLog files with timestamp: ChangeLog_20251210_203502
            if filename.startswith('ChangeLog'):
                print(f"  ✓ Filename pattern: ChangeLog")
                return ('technical', 'documentation', None, [])
            # Tool with hyphen and alphanumeric: gcloud-crc32c, css2
            if re.match(r'^[a-z]+-?[a-z0-9]+$', filename):
                print(f"  ✓ Filename pattern: System utility")
                return ('technical', 'other', None, [])
            # Query parameter style: m=b, m=core, m=RsR2Mc
            if re.match(r'^[a-z]=[a-zA-Z0-9]+$', filename):
                print(f"  ✓ Filename pattern: Query param file")
                return ('technical', 'other', None, [])
            # Makefile patterns: Makefile, Makefile_timestamp
            if filename.startswith('Makefile'):
                print(f"  ✓ Filename pattern: Makefile")
                return ('technical', 'other', None, [])
            # iframe_api, node-which_1 style
            if re.match(r'^[a-z]+[_-][a-z0-9]+(_\d+)?$', filename):
                print(f"  ✓ Filename pattern: System file")
                return ('technical', 'other', None, [])
            # City names with underscore: Rio_Gallegos
            if re.match(r'^[A-Z][a-z]+_[A-Z][a-z]+$', filename):
                print(f"  ✓ Filename pattern: Location data")
                return ('technical', 'other', None, [])
            # Hash with prefix: rs=AA2YrTskOaSug7MVZwlus97OpUaPcMM3bw
            if '=' in filename:
                print(f"  ✓ Filename pattern: Hash parameter file")
                return ('technical', 'other', None, [])
            # Single JS-style file: js(1)
            if re.match(r'^[a-z]+\(\d+\)$', filename):
                print(f"  ✓ Filename pattern: Script copy")
                return ('technical', 'other', None, [])

        # =========================================================
        # PRESENTATION FILES: .pptx, .ppt, .key → Business/Presentations
        # =========================================================
        presentation_extensions = {'.pptx', '.ppt', '.key', '.odp'}
        if ext in presentation_extensions:
            print(f"  ✓ Filename pattern: Presentation ({ext})")
            return ('business', 'presentations', None, [])

        # =========================================================
        # FINANCIAL XLSX FILES: earnings, budget, expenses, revenue
        # =========================================================
        if ext == '.xlsx':
            financial_keywords = ['earnings', 'budget', 'expenses', 'revenue', 'income', 'profit', 'loss', 'financial']
            if any(kw in stem for kw in financial_keywords):
                print(f"  ✓ Filename pattern: Financial spreadsheet")
                return ('financial', 'other', None, [])
            # Data exports: Users.xlsx, Users_timestamp.xlsx (use original stem for case)
            original_stem = file_path.stem
            if re.match(r'^[A-Z][a-z]+s(_\d{8}_\d{6})?$', original_stem):
                print(f"  ✓ Filename pattern: Data export")
                return ('technical', 'data', None, [])

        # =========================================================
        # DOUBLE EXTENSION FILES: something.jpg.jp2
        # =========================================================
        if '.jpg.jp2' in filename.lower() or '.jpeg.jp2' in filename.lower():
            print(f"  ✓ Filename pattern: Converted photo")
            return ('media', 'photos_other', None, [])

        # =========================================================
        # TRAVEL DOCUMENTS: austin_to_bombay, trip_to_paris
        # =========================================================
        if '_to_' in stem or '-to-' in stem:
            print(f"  ✓ Filename pattern: Travel document")
            return ('person', 'travel', None, [])

        # =========================================================
        # ZOUK: Dance events and classes
        # =========================================================
        if 'zouk' in stem:
            print(f"  ✓ Filename pattern: Zouk")
            return ('zouk', 'events', None, [])

        # =========================================================
        # LEGAL/CONTRACT DOCUMENTS: DPA, NDA, SLA, TOS, MSA, SOW
        # Must check BEFORE event documents (dates in filenames)
        # =========================================================
        legal_keywords = ['dpa', 'nda', 'sla', 'tos', 'msa', 'sow', 'contract', 'agreement',
                         'terms', 'privacy', 'policy', 'license', 'eula', 'gdpr', 'hipaa',
                         'compliance', 'legal', 'addendum', 'amendment']
        if ext in {'.pdf', '.docx', '.doc'}:
            if any(kw in stem for kw in legal_keywords):
                print(f"  ✓ Filename pattern: Legal/contract document")
                return ('business', 'legal', None, [])

        # =========================================================
        # EVENT DOCUMENTS: Oct25Event, Nov15Party (month+day in name)
        # =========================================================
        month_patterns = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
        if ext in {'.docx', '.doc', '.pdf'}:
            for month in month_patterns:
                if month in stem and re.search(r'\d{1,2}', stem):
                    print(f"  ✓ Filename pattern: Event document")
                    return ('person', 'events', None, [])

        # =========================================================
        # JOURNAL ENTRIES: Dream, Diary, Thoughts, Reflections
        # =========================================================
        journal_keywords = ['dream', 'diary', 'journal', 'thoughts', 'reflection',
                           'memoir', 'nightbefore', 'morningafter', 'dayof']
        if ext in {'.docx', '.doc', '.txt', '.md'}:
            if any(kw in stem for kw in journal_keywords):
                print(f"  ✓ Filename pattern: Journal entry")
                return ('person', 'other', None, [])

        # =========================================================
        # PERSONAL DOCUMENTS: Short name + version (Sumedh3.docx)
        # =========================================================
        if ext in {'.docx', '.doc'}:
            original_stem = file_path.stem
            # Short name followed by digit (personal documents)
            if re.match(r'^[A-Z][a-z]+\d$', original_stem):
                print(f"  ✓ Filename pattern: Personal document")
                return ('person', 'other', None, [])
            # PascalCase event names (ZoukSocial, DanceNight)
            if re.match(r'^([A-Z][a-z]+){2,}$', original_stem):
                print(f"  ✓ Filename pattern: Event document")
                return ('person', 'events', None, [])

        # =========================================================
        # PITCH/PROPOSAL FILES: Pitch, Proposal with version
        # =========================================================
        if ext in {'.pptx', '.pdf', '.docx'}:
            if stem.startswith('pitch') or stem.startswith('proposal'):
                print(f"  ✓ Filename pattern: Business pitch/proposal")
                return ('business', 'presentations', None, [])

        return None

    def extract_text_from_image(self, image_path: Path) -> str:
        """Extract text from image using OCR."""
        if not self.ocr_available:
            return ""

        with CostTracker(self.cost_calculator, 'tesseract_ocr') if self.cost_calculator else nullcontext():
            try:
                image = Image.open(image_path)
                text = pytesseract.image_to_string(image)
                return text.strip()
            except Exception as e:
                print(f"  OCR error: {e}")
                return ""

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF (searchable or scanned)."""
        if not self.ocr_available:
            return ""

        with CostTracker(self.cost_calculator, 'pdf_extraction') if self.cost_calculator else nullcontext():
            text = ""

            try:
                # First try to extract text directly (for searchable PDFs)
                with open(pdf_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages[:10]:  # Limit to first 10 pages
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"

                # If we got meaningful text, return it
                if len(text.strip()) > 100:
                    return text.strip()

                # Otherwise, try OCR on the PDF
                print(f"  Using OCR for scanned PDF...")
                images = convert_from_path(pdf_path, first_page=1, last_page=5)
                for image in images:
                    text += pytesseract.image_to_string(image) + "\n"

                return text.strip()
            except Exception as e:
                print(f"  PDF extraction error: {e}")
                return ""

    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            return ""

        with CostTracker(self.cost_calculator, 'docx_extraction') if self.cost_calculator else nullcontext():
            try:
                doc = Document(docx_path)
                text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text.append(paragraph.text)

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text.append(cell.text)

                return "\n".join(text)
            except Exception as e:
                print(f"  DOCX extraction error: {e}")
                return ""

    def extract_text_from_xlsx(self, xlsx_path: Path) -> str:
        """Extract text from Excel spreadsheet."""
        if not EXCEL_AVAILABLE:
            return ""

        with CostTracker(self.cost_calculator, 'xlsx_extraction') if self.cost_calculator else nullcontext():
            try:
                workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
                text = []

                # Limit to first 5 sheets
                for sheet_name in list(workbook.sheetnames)[:5]:
                    sheet = workbook[sheet_name]
                    # Limit to first 100 rows
                    for row in list(sheet.iter_rows(max_row=100, values_only=True)):
                        row_text = ' '.join([str(cell) for cell in row if cell is not None])
                        if row_text.strip():
                            text.append(row_text)

                workbook.close()
                return "\n".join(text)
            except Exception as e:
                print(f"  XLSX extraction error: {e}")
                return ""

    def extract_text(self, file_path: Path) -> str:
        """Extract text from various file types."""
        mime_type = self.enricher.detect_mime_type(str(file_path))
        file_ext = file_path.suffix.lower()

        # Images
        if mime_type and mime_type.startswith('image/'):
            return self.extract_text_from_image(file_path)

        # PDFs
        elif mime_type == 'application/pdf' or file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)

        # Word documents
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)

        # Excel spreadsheets
        elif file_ext in ['.xlsx', '.xls']:
            return self.extract_text_from_xlsx(file_path)

        # Text files
        elif mime_type and mime_type.startswith('text/') or file_ext in ['.txt', '.md', '.csv']:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(50000)  # First 50KB
            except Exception:
                return ""

        return ""

    _GEOGRAPHIC_LABELS = frozenset({
        "a landscape or nature scene",
        "a cityscape or urban scene",
        "a building or architecture",
    })

    def _map_clip_label(self, label: str, image_metadata: Dict = None) -> Optional[Tuple[str, str]]:
        """Map a CLIP label to (category, subcategory), upgrading to travel if GPS present."""
        mapping = CLIP_LABEL_TO_ORGANIZER.get(label)
        if not mapping:
            return None
        cat, subcat = mapping
        if image_metadata and image_metadata.get("gps_coordinates") and label in self._GEOGRAPHIC_LABELS:
            cat, subcat = "media", "photos_travel"
        return (cat, subcat)

    def enhance_weak_image_classification(
        self, file_path: Path, image_metadata: Dict = None
    ) -> Optional[Tuple[str, str]]:
        """Run full 20-category CLIP + OCR fallback for weakly classified images.

        Only called for images that would otherwise land in photos_other or uncategorized.
        Returns (category, subcategory) or None to keep original classification.
        """
        if not ENHANCED_CLIP_AVAILABLE or not self.image_analyzer.vision_available:
            return None
        if not self.image_analyzer.model or not self.image_analyzer.processor:
            return None

        try:
            image = Image.open(file_path)
        except Exception as e:
            print(f"  CLIP enhance: cannot open image: {e}")
            return None

        # Run full 20-category CLIP classification
        try:
            inputs = self.image_analyzer.processor(
                text=CLIP_CATEGORY_PROMPTS, images=image,
                return_tensors="pt", padding=True,
            )
            with torch.no_grad():
                probs = self.image_analyzer.model(**inputs).logits_per_image.softmax(dim=1)
            scores = {label: float(probs[0][i]) for i, label in enumerate(CLIP_CONTENT_LABELS)}
            best_label = max(scores, key=scores.get)
            best_score = scores[best_label]
            print(f"  CLIP enhance: {best_label} ({best_score:.1%})")
        except Exception as e:
            print(f"  CLIP enhance error: {e}")
            return None

        if best_score < CLIP_ENHANCE_THRESHOLD:
            return None

        # High confidence — map directly, skip OCR
        if best_score >= CLIP_ENHANCE_HIGH_THRESHOLD:
            result = self._map_clip_label(best_label, image_metadata)
            if result:
                print(f"  CLIP enhance → {result[0]}/{result[1]} (high confidence)")
                return result

        # Medium confidence — try OCR first, fall back to CLIP mapping
        if self.ocr_available:
            try:
                ocr_text = self.extract_text_from_image(file_path)
                if ocr_text and len(ocr_text) >= 30:
                    text_cat, text_subcat, _, _ = self.classifier.classify_content(ocr_text, file_path.name)
                    if text_cat != "uncategorized":
                        print(f"  CLIP enhance → {text_cat}/{text_subcat} (OCR fallback)")
                        return (text_cat, text_subcat)
            except Exception as e:
                print(f"  CLIP enhance OCR error: {e}")

        result = self._map_clip_label(best_label, image_metadata)
        if result:
            print(f"  CLIP enhance → {result[0]}/{result[1]} (medium confidence)")
        return result

    def detect_file_category(self, file_path: Path) -> Tuple[str, str, str, str, Optional[str], List[str], Dict[str, Any]]:
        """
        Detect file category based on content.

        Priority order:
        0. Filename pattern detection (fastest - no content extraction needed)
        1. Organization entity detection (for documents with content)
        2. Person entity detection
        3. Game asset detection (audio, sprites, textures)
        4. Filepath-based classification (file extensions, filenames)
        5. Image content analysis (for home interiors)
        6. OCR and text-based classification

        Returns:
            Tuple of (main_category, subcategory, schema_type, extracted_text, company_name, people_names, image_metadata)
        """
        # Determine schema type and MIME type early (needed for multiple paths)
        mime_type = self.enricher.detect_mime_type(str(file_path))
        if mime_type:
            if mime_type.startswith('image/'):
                schema_type = 'ImageObject'
            elif mime_type == 'application/pdf':
                schema_type = 'DigitalDocument'
            elif mime_type.startswith('video/'):
                schema_type = 'VideoObject'
            elif mime_type.startswith('audio/'):
                schema_type = 'AudioObject'
            else:
                schema_type = 'DigitalDocument'
        else:
            schema_type = 'DigitalDocument'

        # PRIORITY 0: Filename pattern detection (fastest - no content extraction needed)
        # Handles: Google invoices, resumes, technical files, legal docs, business docs, entity files
        filename_result = self.classify_by_filename_patterns(file_path)
        if filename_result:
            category, subcategory, company_name, people_names = filename_result
            # Handle skip category for duplicates
            if category == 'skip':
                return ('skip', subcategory, schema_type, '', None, [], {})
            # Point A: enhance weak photos_other from filename patterns for images
            if subcategory == 'photos_other' and schema_type == 'ImageObject':
                enhanced = self.enhance_weak_image_classification(file_path)
                if enhanced:
                    return (enhanced[0], enhanced[1], schema_type, '', None, [], {})
            return (category, subcategory, schema_type, '', company_name, people_names, {})

        # PRIORITY 1: Organization and Person detection for document-type files
        # Only apply to document/PDF types (not images, audio, video)
        if schema_type == 'DigitalDocument' or mime_type == 'application/pdf':
            print(f"  Checking for Organization/Person entities...")
            extracted_text = self.extract_text(file_path)

            if extracted_text and len(extracted_text) >= 50:
                # Try Organization detection first
                org_result = self.classify_by_organization(extracted_text, file_path.name)
                if org_result:
                    category, subcategory, org_name = org_result
                    print(f"  ✓ Organization detected: {org_name} ({subcategory})")
                    return (category, subcategory, schema_type, extracted_text, org_name, [], {})

                # Try Person detection second
                person_result = self.classify_by_person(extracted_text, file_path.name)
                if person_result:
                    category, subcategory, people_names = person_result
                    print(f"  ✓ Person detected: {', '.join(people_names[:3]) if people_names else 'Unknown'} ({subcategory})")
                    return (category, subcategory, schema_type, extracted_text, None, people_names, {})

        # PRIORITY 3: Check for game assets (before filepath patterns)
        game_asset = self.classify_game_asset(file_path)
        if game_asset:
            category, subcategory = game_asset
            print(f"  ✓ Game asset detected: {subcategory}")
            return (category, subcategory, schema_type, '', None, [], {})

        # PRIORITY 3: Check filepath patterns (most efficient and accurate for code files)
        filepath_category = self.classify_by_filepath(file_path)
        if filepath_category:
            print(f"  ✓ Filepath match: {filepath_category}")
            # Return filepath-based category as a special marker
            # We'll handle this in get_destination_path
            return ('filepath', filepath_category, schema_type, '', None, [], {})

        # Extract metadata for images
        image_metadata = {}
        if schema_type == 'ImageObject' and self.metadata_parser.metadata_available:
            print(f"  Extracting image metadata...")
            image_metadata = self.metadata_parser.get_metadata_summary(file_path)

            if image_metadata.get('datetime'):
                dt = image_metadata['datetime']
                print(f"  ✓ Photo taken: {dt.strftime('%Y-%m-%d %H:%M:%S')}")

            if image_metadata.get('gps_coordinates'):
                coords = image_metadata['gps_coordinates']
                print(f"  ✓ GPS: {coords[0]:.6f}, {coords[1]:.6f}")

            if image_metadata.get('location_name'):
                print(f"  ✓ Location: {image_metadata['location_name']}")

        # PRIORITY 3.5: Check for identification documents in images (passport, ID, license)
        # These should go to Person/ folder, not Media/
        if schema_type == 'ImageObject' and self.ocr_available:
            # Extract text from image via OCR
            ocr_text = self.extract_text_from_image(file_path)
            if ocr_text and len(ocr_text) >= 30:
                ocr_lower = ocr_text.lower()
                # Check for identification document keywords
                id_keywords = ['passport', 'driver license', "driver's license", 'identification',
                              'united states of america', 'department of state', 'nationality',
                              'date of birth', 'place of birth', 'surname', 'given names',
                              'social security', 'state id', 'national id']
                if any(kw in ocr_lower for kw in id_keywords):
                    print(f"  ✓ Identification document detected via OCR")
                    people_names = []

                    # Method 1: Parse passport MRZ (Machine Readable Zone)
                    # Format: P<COUNTRY{SURNAME}<<{GIVEN_NAME}<...
                    mrz_match = re.search(r'P<[A-Z]{3}([A-Z]+)<<([A-Z]+)<', ocr_text)
                    if mrz_match:
                        surname = mrz_match.group(1).title()
                        given = mrz_match.group(2).title()
                        people_names = [f"{given} {surname}"]

                    # Method 2: Look for name fields with values on next line or after colon
                    # Passport format: "Surname\nLEDLIE" or "Surname/Nom\nLEDLIE"
                    if not people_names:
                        # Find surname (all caps, standalone on line)
                        surname_match = re.search(r'(?:surname|nom|apellidos)[/\w\s]*\n\s*([A-Z]{2,})\b', ocr_text, re.IGNORECASE)
                        given_match = re.search(r'(?:given\s*names?|pr[ée]noms?|nombres)[/\w\s]*\n\s*([A-Z]{2,})\b', ocr_text, re.IGNORECASE)
                        if surname_match and given_match:
                            people_names = [f"{given_match.group(1).title()} {surname_match.group(1).title()}"]

                    # Method 3: General name extraction patterns
                    if not people_names:
                        people_names = self.classifier.extract_people_names(ocr_text)

                    if people_names:
                        print(f"  ✓ Person identified: {people_names[0]}")
                    return ('person', 'contacts', schema_type, ocr_text, None, people_names, image_metadata)

        # PRIORITY 4: Check for media files (photos, videos, audio)
        # This runs after metadata extraction so we can use GPS/datetime for classification
        media_classification = self.classify_media_file(file_path, image_metadata)
        if media_classification:
            category, media_type, subcategory = media_classification
            # Point B: enhance weak photos/other for images
            if media_type == 'photos' and subcategory == 'other':
                enhanced = self.enhance_weak_image_classification(file_path, image_metadata)
                if enhanced:
                    print(f"  ✓ Enhanced media: {enhanced[0]}/{enhanced[1]}")
                    return (enhanced[0], enhanced[1], schema_type, '', None, [], image_metadata)
            print(f"  ✓ Media file detected: {media_type}/{subcategory}")
            return (category, f"{media_type}_{subcategory}", schema_type, '', None, [], image_metadata)

        # PRIORITY 5: Check for photos with people (social)
        if schema_type == 'ImageObject' and self.image_analyzer.vision_available:
            print(f"  Analyzing image content...")

            # First check if photo has people
            has_people, people_scores = self.image_analyzer.has_people_in_photo(file_path)

            if has_people:
                print(f"  ✓ Detected: Photo with people")
                if people_scores:
                    top_categories = sorted(people_scores.items(), key=lambda x: x[1], reverse=True)[:3]
                    print(f"  Top matches: {', '.join([f'{cat}: {score:.2%}' for cat, score in top_categories])}")
                return ('media', 'photos_social', schema_type, '', None, [], image_metadata)

            # Then check for home interior without people
            is_property_mgmt, vision_scores = self.image_analyzer.is_home_interior_no_people(file_path)

            if is_property_mgmt:
                print(f"  ✓ Detected: Home interior without people")
                if vision_scores:
                    top_categories = sorted(vision_scores.items(), key=lambda x: x[1], reverse=True)[:3]
                    print(f"  Top matches: {', '.join([f'{cat}: {score:.2%}' for cat, score in top_categories])}")
                return ('property_management', 'other', schema_type, '', None, [], image_metadata)

        # PRIORITY 6: Regular text extraction and classification
        print(f"  Extracting content...")
        extracted_text = self.extract_text(file_path)

        if extracted_text:
            print(f"  Extracted {len(extracted_text)} characters")
            category, subcategory, company_name, people_names = self.classifier.classify_content(extracted_text, file_path.name)
            if company_name:
                print(f"  Detected company: {company_name}")
            if people_names:
                print(f"  Detected people: {', '.join(people_names[:3])}{' ...' if len(people_names) > 3 else ''}")
            print(f"  Classified as: {category}/{subcategory}")
        else:
            print(f"  No text extracted, using filename")
            category, subcategory, company_name, people_names = self.classifier.classify_content("", file_path.name)

        # Point C: last-resort enhancement for uncategorized images
        if category == 'uncategorized' and schema_type == 'ImageObject':
            enhanced = self.enhance_weak_image_classification(file_path, image_metadata)
            if enhanced:
                print(f"  ✓ Enhanced uncategorized: {enhanced[0]}/{enhanced[1]}")
                return (enhanced[0], enhanced[1], schema_type, extracted_text, None, [], image_metadata)

        return (category, subcategory, schema_type, extracted_text, company_name, people_names, image_metadata)

    def generate_schema(self, file_path: Path, schema_type: str, extracted_text: str = "") -> Dict:
        """Generate Schema.org metadata for a file with extracted content."""
        stats = file_path.stat()
        mime_type = self.enricher.detect_mime_type(str(file_path))
        file_url = f"https://localhost/files/{quote(file_path.name)}"
        actual_path = str(file_path.absolute())

        # Create generator based on type
        if schema_type == 'ImageObject':
            generator = ImageGenerator(schema_type)
            generator.set_basic_info(
                name=file_path.name,
                content_url=file_url,
                encoding_format=mime_type or 'image/png',
                description=f"{file_path.name}"
            )
        elif schema_type in ['DigitalDocument', 'Article']:
            generator = DocumentGenerator(schema_type)
            generator.set_basic_info(
                name=file_path.name,
                description=f"{file_path.name}"
            )
            generator.set_file_info(
                encoding_format=mime_type or 'application/octet-stream',
                url=file_url,
                content_size=stats.st_size
            )
        else:
            generator = DocumentGenerator()
            generator.set_basic_info(
                name=file_path.name,
                description=f"{file_path.name}"
            )

        # Set dates
        try:
            generator.set_dates(
                created=datetime.fromtimestamp(stats.st_ctime),
                modified=datetime.fromtimestamp(stats.st_mtime)
            )
        except Exception:
            pass

        # Add extracted text as abstract/text property
        if extracted_text:
            try:
                # Truncate to reasonable length for schema
                text_preview = extracted_text[:1000] + ('...' if len(extracted_text) > 1000 else '')
                generator.set_property('abstract', text_preview, PropertyType.TEXT)
                generator.set_property('text', extracted_text[:5000], PropertyType.TEXT)
            except Exception:
                pass

        # Add file path
        try:
            generator.set_property('filePath', actual_path, PropertyType.TEXT)
        except Exception:
            pass

        return generator.to_dict()

    def get_destination_path(self, file_path: Path, category: str, subcategory: str, company_name: Optional[str] = None, image_metadata: Optional[Dict] = None, people_names: Optional[List[str]] = None) -> Path:
        """
        Get the destination path for a file based on content category.

        Args:
            file_path: Path to the file
            category: Main category
            subcategory: Subcategory
            company_name: Optional company name for business/organization files
            image_metadata: Optional metadata for images (datetime, location)
            people_names: Optional list of people names for person-classified files

        Returns:
            Destination path for the file
        """
        # Special handling for filepath-based classification
        if category == 'filepath':
            # subcategory contains the full path (e.g., "Technical/Python/MyProject")
            relative_path = subcategory
        # Special handling for media files with nested structure
        elif category == 'media' and '_' in subcategory:
            # subcategory format: "photos_screenshots" or "photos_screenshots_browser"
            parts = subcategory.split('_', 1)  # Split into at most 2 parts
            if len(parts) == 2:
                media_type, media_subcat = parts
                if media_type in self.category_paths['media']:
                    media_dict = self.category_paths['media'][media_type]
                    if isinstance(media_dict, dict):
                        # Check for 3-level nesting (e.g., screenshots_browser)
                        if '_' in media_subcat:
                            parent_key, child_key = media_subcat.split('_', 1)
                            parent_val = media_dict.get(parent_key)
                            if isinstance(parent_val, dict):
                                relative_path = parent_val.get(child_key, parent_val.get('other', f'Media/{media_type.capitalize()}/{parent_key.capitalize()}'))
                            else:
                                relative_path = media_dict.get(media_subcat, media_dict.get('other', f'Media/{media_type.capitalize()}/Other'))
                        else:
                            val = media_dict.get(media_subcat)
                            if isinstance(val, dict):
                                relative_path = val.get('other', f'Media/{media_type.capitalize()}/{media_subcat.capitalize()}')
                            elif val:
                                relative_path = val
                            else:
                                relative_path = media_dict.get('other', f'Media/{media_type.capitalize()}/Other')
                    else:
                        relative_path = media_dict
                else:
                    relative_path = 'Media/Other'
            else:
                relative_path = 'Media/Other'
        elif category in self.category_paths:
            if isinstance(self.category_paths[category], dict):
                if subcategory in self.category_paths[category]:
                    relative_path = self.category_paths[category][subcategory]
                else:
                    relative_path = self.category_paths[category].get('other', f'{category.capitalize()}/Other')
            else:
                relative_path = self.category_paths[category]
        else:
            relative_path = 'Uncategorized'

        # Organization: Create entity-named subfolders under Organization/
        # Structure: Organization/{OrgName}/ for most types
        # Exception: Organization/Clients/{OrgName}/ for clients (nested subfolders)
        if category == 'organization' and company_name:
            sanitized_company = self.classifier.sanitize_company_name(company_name)
            # Only create company subfolder if name is valid (not a sentence fragment)
            if sanitized_company:
                if subcategory == 'clients':
                    # Clients get nested: Organization/Clients/{OrgName}/
                    relative_path = f"{relative_path}/{sanitized_company}"
                elif subcategory == 'meeting_notes':
                    # Meeting notes get nested: Organization/{OrgName}/Meeting Notes/
                    relative_path = f"{relative_path}/{sanitized_company}/Meeting Notes"
                else:
                    # All other org types: Organization/{OrgName}/
                    relative_path = f"{relative_path}/{sanitized_company}"

        # Person: Create person-named subfolders under Person/
        # Structure: Person/{PersonName}/ for all types
        if category == 'person' and people_names:
            # Use first person name as the folder name
            person_name = people_names[0] if people_names else 'Unknown'
            sanitized_person = self.classifier.sanitize_company_name(person_name)
            # Only create person subfolder if name is valid
            if sanitized_person:
                relative_path = f"{relative_path}/{sanitized_person}"
            else:
                relative_path = f"{relative_path}/Unknown"
        elif category == 'person' and not people_names:
            # Fallback for person category without extracted names
            relative_path = f"{relative_path}/Unknown"

        # Legacy: client files from business category with company name
        if category == 'business' and subcategory == 'clients' and company_name:
            sanitized_company = self.classifier.sanitize_company_name(company_name)
            # Only create company subfolder if name is valid
            if sanitized_company:
                relative_path = f"{relative_path}/{sanitized_company}"

        # Date-based organization for images (if enabled and metadata available)
        if self.organize_by_date and image_metadata and image_metadata.get('year'):
            year = image_metadata['year']
            month = image_metadata['month']
            relative_path = f"Photos/{year}/{month:02d}"

        # Location-based organization for images (if enabled and location available)
        elif self.organize_by_location and image_metadata and image_metadata.get('location_name'):
            # Clean location name for folder
            location = image_metadata['location_name']
            # Take first part (usually city)
            city = location.split(',')[0].strip()
            # Sanitize for folder name
            safe_city = re.sub(r'[<>:"/\\|?*]', '', city)
            relative_path = f"Photos/Locations/{safe_city}"

        dest_dir = self.base_path / relative_path
        dest_dir.mkdir(parents=True, exist_ok=True)

        # Handle duplicate filenames
        dest_path = dest_dir / file_path.name
        if dest_path.exists() and dest_path != file_path:
            stem = file_path.stem
            suffix = file_path.suffix
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            dest_path = dest_dir / f"{stem}_{timestamp}{suffix}"

        return dest_path

    def should_skip_file(self, file_path: Path) -> bool:
        """Check if file should be skipped."""
        skip_files = {'.DS_Store', '.localized', 'Thumbs.db', 'desktop.ini'}
        skip_dirs = {'__pycache__', '.git', 'node_modules', '.venv', 'venv'}

        if file_path.name.startswith('.') and file_path.name not in {'.gitignore', '.env.example'}:
            return True

        if file_path.name in skip_files:
            return True

        if any(skip_dir in file_path.parts for skip_dir in skip_dirs):
            return True

        return False

    def _persist_to_graph_store(
        self,
        file_path: Path,
        dest_path: Path,
        category: str,
        subcategory: str,
        schema: Dict,
        extracted_text: str,
        company_name: Optional[str],
        people_names: List[str],
        image_metadata: Optional[Dict]
    ) -> None:
        """
        Persist file and its relationships to the graph store with canonical IDs.

        This method creates:
        - File record with canonical_id (urn:sha256:{hash})
        - Category record with canonical_id (UUID v5 from name)
        - Company record with canonical_id (UUID v5 from name)
        - Person records with canonical_id (UUID v5 from name)
        - Location record with canonical_id (UUID v5 from name)
        - Relationships between file and entities
        """
        try:
            session = self.graph_store.get_session()

            # Get file stats
            stat = file_path.stat() if file_path.exists() else dest_path.stat()

            # Add file to store (generates canonical_id automatically)
            file_record = self.graph_store.add_file(
                original_path=str(file_path),
                filename=file_path.name,
                session=session,
                current_path=str(dest_path),
                file_size=stat.st_size,
                mime_type=schema.get('encodingFormat'),
                schema_type=schema.get('@type'),
                schema_data=schema,
                extracted_text=extracted_text[:10000] if extracted_text else None,
                extracted_text_length=len(extracted_text) if extracted_text else 0,
                status=FileStatus.ORGANIZED,
                organized_at=datetime.now()
            )

            file_id = file_record.id

            # Add category relationship
            self.graph_store.add_file_to_category(
                file_id=file_id,
                category_name=category,
                subcategory_name=subcategory,
                session=session
            )

            # Add company relationship if detected
            if company_name:
                self.graph_store.add_file_to_company(
                    file_id=file_id,
                    company_name=company_name,
                    context='content_analysis',
                    session=session
                )

            # Add people relationships if detected
            if people_names:
                for person_name in people_names:
                    self.graph_store.add_file_to_person(
                        file_id=file_id,
                        person_name=person_name,
                        role='mentioned',
                        session=session
                    )

            # Add location if available from image metadata
            if image_metadata and image_metadata.get('location'):
                location_info = image_metadata['location']
                self.graph_store.add_file_to_location(
                    file_id=file_id,
                    location_name=location_info.get('display_name', 'Unknown'),
                    latitude=location_info.get('latitude'),
                    longitude=location_info.get('longitude'),
                    city=location_info.get('city'),
                    state=location_info.get('state'),
                    country=location_info.get('country'),
                    location_type='captured_at',
                    session=session
                )

            session.commit()
            session.close()

        except Exception as e:
            print(f"  ⚠ Graph store error (non-fatal): {e}")

    def organize_file(self, file_path: Path, dry_run: bool = False, force: bool = False) -> Dict:
        """
        Organize a single file based on content.

        Args:
            file_path: Path to the file
            dry_run: If True, don't actually move files
            force: If True, re-organize even if already in correct location

        Returns:
            Dictionary with organization details
        """
        result = {
            'source': str(file_path),
            'status': 'skipped',
            'reason': None,
            'destination': None,
            'schema': None,
            'extracted_text_length': 0
        }

        if self.should_skip_file(file_path):
            result['reason'] = 'system_file'
            self.stats['skipped'] += 1
            return result

        if not file_path.is_file():
            result['reason'] = 'not_file'
            self.stats['skipped'] += 1
            return result

        try:
            # Detect category based on content
            category, subcategory, schema_type, extracted_text, company_name, people_names, image_metadata = self.detect_file_category(file_path)
            result['extracted_text_length'] = len(extracted_text)
            result['company_name'] = company_name
            result['people_names'] = people_names
            result['image_metadata'] = image_metadata

            # Handle skip category (duplicates, etc.)
            if category == 'skip':
                result['status'] = 'skipped'
                result['reason'] = subcategory  # e.g., 'duplicate'
                self.stats['skipped'] += 1
                return result

            # Generate schema with extracted content
            schema = self.generate_schema(file_path, schema_type, extracted_text)

            # Validate schema
            validation_report = self.validator.validate(schema)

            # Get destination path (with optional date/location organization for images)
            dest_path = self.get_destination_path(file_path, category, subcategory, company_name, image_metadata, people_names)

            # Skip if already in the right place (unless force=True)
            if file_path == dest_path and not force:
                result['status'] = 'already_organized'
                result['destination'] = str(dest_path)
                result['schema'] = schema
                result['category'] = category
                result['subcategory'] = subcategory
                self.stats['already_organized'] += 1
                return result

            # Move file if not dry run
            if not dry_run:
                shutil.move(str(file_path), str(dest_path))

                # Register schema
                schema['url'] = f"file://{dest_path.absolute()}"
                metadata = {
                    'category': category,
                    'subcategory': subcategory,
                    'organized_date': datetime.now().isoformat(),
                    'is_valid': validation_report.is_valid(),
                    'has_extracted_text': bool(extracted_text)
                }
                if company_name:
                    metadata['company_name'] = company_name

                self.registry.register(
                    str(dest_path),
                    schema,
                    metadata=metadata
                )

                # Persist to database with canonical IDs
                if self.graph_store:
                    self._persist_to_graph_store(
                        file_path=file_path,
                        dest_path=dest_path,
                        category=category,
                        subcategory=subcategory,
                        schema=schema,
                        extracted_text=extracted_text,
                        company_name=company_name,
                        people_names=people_names,
                        image_metadata=image_metadata
                    )

            result['status'] = 'organized' if not dry_run else 'would_organize'
            result['destination'] = str(dest_path)
            result['schema'] = schema
            result['category'] = category
            result['subcategory'] = subcategory
            result['is_valid'] = validation_report.is_valid()

            self.stats['organized'] += 1
            self.stats[f'{category}_{subcategory}'] += 1

        except Exception as e:
            result['status'] = 'error'
            result['reason'] = str(e)
            self.stats['errors'] += 1
            print(f"  ✗ Error: {e}")

        return result

    def scan_directory(self, directory: Path) -> List[Path]:
        """Scan directory for files to organize."""
        files = []
        try:
            for item in directory.rglob('*'):
                if item.is_file() and not self.should_skip_file(item):
                    files.append(item)
        except PermissionError:
            print(f"Permission denied: {directory}")
        return files

    def organize_directories(self, source_dirs: List[str], dry_run: bool = False, limit: int = None, force: bool = False) -> Dict:
        """
        Organize files from multiple source directories.

        Args:
            source_dirs: List of directory paths to organize
            dry_run: If True, simulate organization without moving files
            limit: Maximum number of files to process (for testing)
            force: If True, re-organize files even if already in correct location

        Returns:
            Dictionary with organization results
        """
        results = []

        print(f"\n{'='*60}")
        print(f"Content-Based File Organization {'(DRY RUN)' if dry_run else ''}")
        print(f"{'='*60}\n")

        if not self.ocr_available:
            print("⚠️  WARNING: OCR libraries not available")
            print("   Install with: pip install pytesseract Pillow pypdf pdf2image")
            print("   Content classification will be limited to filenames\n")

        # Scan all directories
        all_files = []
        for source_dir in source_dirs:
            source_path = Path(source_dir).expanduser()
            if source_path.exists():
                print(f"Scanning: {source_path}")
                files = self.scan_directory(source_path)
                all_files.extend(files)
                print(f"  Found {len(files)} files")
            else:
                print(f"Directory not found: {source_path}")

        if limit:
            all_files = all_files[:limit]
            print(f"\n⚠️  Processing limited to first {limit} files for testing\n")

        print(f"\nTotal files to process: {len(all_files)}\n")

        # Organize each file
        for i, file_path in enumerate(all_files, 1):
            print(f"[{i}/{len(all_files)}] Processing: {file_path.name}")
            result = self.organize_file(file_path, dry_run=dry_run, force=force)
            results.append(result)

            if result['status'] == 'organized' or result['status'] == 'would_organize':
                print(f"  → {result['destination']}")
            elif result['status'] == 'error':
                print(f"  ✗ Error: {result['reason']}")

        # Generate summary
        summary = {
            'total_files': len(all_files),
            'organized': self.stats['organized'],
            'already_organized': self.stats['already_organized'],
            'skipped': self.stats['skipped'],
            'errors': self.stats['errors'],
            'dry_run': dry_run,
            'results': results,
            'registry_stats': self.registry.get_statistics() if not dry_run else None
        }

        return summary

    def print_summary(self, summary: Dict):
        """Print organization summary."""
        print(f"\n{'='*60}")
        print("Organization Summary")
        print(f"{'='*60}\n")

        print(f"Total files processed: {summary['total_files']}")
        print(f"Successfully organized: {summary['organized']}")
        print(f"Already organized: {summary['already_organized']}")
        print(f"Skipped: {summary['skipped']}")
        print(f"Errors: {summary['errors']}")

        if summary['dry_run']:
            print("\n⚠️  This was a DRY RUN - no files were moved")

        # Category breakdown
        print(f"\n{'='*60}")
        print("Category Breakdown")
        print(f"{'='*60}\n")

        category_stats = defaultdict(int)
        for result in summary['results']:
            if result.get('category'):
                category_stats[result['category']] += 1

        for category, count in sorted(category_stats.items()):
            print(f"{category.capitalize()}: {count} files")

        # OCR stats
        ocr_count = sum(1 for r in summary['results'] if r.get('extracted_text_length', 0) > 0)
        print(f"\n{'='*60}")
        print("Content Extraction Stats")
        print(f"{'='*60}\n")
        print(f"Files with extracted text: {ocr_count}/{summary['total_files']}")

        # Company detection stats
        company_files = [r for r in summary['results'] if r.get('company_name')]
        if company_files:
            print(f"\n{'='*60}")
            print("Detected Companies")
            print(f"{'='*60}\n")
            company_counts = defaultdict(int)
            for result in company_files:
                company_counts[result['company_name']] += 1

            print(f"Total files with detected companies: {len(company_files)}")
            print(f"\nCompanies found:")
            for company, count in sorted(company_counts.items(), key=lambda x: x[1], reverse=True):
                print(f"  {company}: {count} files")

        if summary.get('registry_stats'):
            print(f"\n{'='*60}")
            print("Schema Registry")
            print(f"{'='*60}\n")
            stats = summary['registry_stats']
            print(f"Total schemas: {stats['total_schemas']}")
            print(f"Types: {', '.join(stats['types'])}")

        # Cost tracking summary
        if self.cost_calculator:
            self._print_cost_summary()

    def _print_cost_summary(self):
        """Print cost and ROI summary from the cost calculator."""
        if not self.cost_calculator:
            return

        print(f"\n{'='*60}")
        print("Cost & ROI Analysis")
        print(f"{'='*60}\n")

        cost_summary = self.cost_calculator.calculate_total_cost()
        roi_summary = self.cost_calculator.calculate_total_roi()

        print(f"Total Processing Cost:     ${cost_summary['total_cost']:.4f}")
        print(f"Total Files Processed:     {cost_summary['total_files_processed']:,}")
        print(f"Avg Cost per File:         ${cost_summary['avg_cost_per_file']:.6f}")
        print(f"Total Processing Time:     {cost_summary['total_processing_time_sec']:.1f}s")

        print(f"\nEstimated Value Generated: ${roi_summary['total_value']:.2f}")
        roi_pct = roi_summary['overall_roi_percentage']
        roi_str = f"{roi_pct:.0f}%" if roi_pct != float('inf') else "∞"
        print(f"Overall ROI:               {roi_str}")
        print(f"Manual Hours Saved:        {roi_summary['total_manual_hours_saved']:.1f} hours")

        # Per-feature breakdown (top 5 by usage)
        feature_costs = cost_summary.get('feature_breakdown', {})
        if feature_costs:
            print(f"\n{'Feature':<25} {'Cost':>10} {'Files':>10}")
            print("-" * 50)
            sorted_features = sorted(
                feature_costs.items(),
                key=lambda x: x[1]['total_files_processed'],
                reverse=True
            )
            for feature_name, data in sorted_features[:7]:
                if data['total_invocations'] > 0:
                    print(f"{feature_name:<25} ${data['total_cost']:>9.4f} {data['total_files_processed']:>10,}")

        # Show recommendations if any critical issues
        recommendations = self.cost_calculator.get_optimization_recommendations()
        critical_recs = [r for r in recommendations if r['severity'] in ('critical', 'high')]
        if critical_recs:
            print(f"\n⚠️  Optimization Recommendations:")
            for rec in critical_recs[:3]:
                print(f"   • {rec['message']}")

    def get_cost_report(self) -> Optional[Dict[str, Any]]:
        """
        Get the full cost and ROI report.

        Returns:
            Cost report dictionary or None if cost tracking is disabled
        """
        if not self.cost_calculator:
            return None
        return self.cost_calculator.generate_report()

    def save_cost_report(self, output_path: str = None):
        """
        Save the cost report to a JSON file.

        Args:
            output_path: Path to save the report (auto-generated if None)
        """
        if not self.cost_calculator:
            print("Cost tracking is not enabled")
            return

        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"results/cost_report_{timestamp}.json"

        self.cost_calculator.generate_report(output_path)
        print(f"Cost report saved to: {output_path}")

    def save_report(self, summary: Dict, output_path: str = None):
        """Save detailed organization report to JSON."""
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"results/content_organization_report_{timestamp}.json"

        with open(output_path, 'w') as f:
            json.dump(summary, f, indent=2, default=str)

        print(f"\nDetailed report saved to: {output_path}")


def main():
    """Main entry point."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Organize files by content using OCR and Schema.org metadata'
    )
    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Simulate organization without moving files'
    )
    parser.add_argument(
        '--base-path',
        default='~/Documents',
        help='Base path for organized files (default: ~/Documents)'
    )
    parser.add_argument(
        '--sources',
        nargs='+',
        default=['~/Desktop', '~/Downloads'],
        help='Source directories to organize (default: ~/Desktop ~/Downloads)'
    )
    parser.add_argument(
        '--report',
        help='Path to save detailed JSON report'
    )
    parser.add_argument(
        '--limit',
        type=int,
        help='Limit number of files to process (for testing)'
    )
    parser.add_argument(
        '--no-cost-tracking',
        action='store_true',
        help='Disable cost and ROI tracking'
    )
    parser.add_argument(
        '--cost-report',
        nargs='?',
        const='results/cost_report.json',
        default='results/cost_report.json',
        help='Path to save cost/ROI report (default: results/cost_report.json, use --no-cost-tracking to disable)'
    )
    parser.add_argument(
        '--check-deps',
        action='store_true',
        help='Run system health check and show feature availability'
    )
    parser.add_argument(
        '--skip-health-check',
        action='store_true',
        help='Skip startup health check'
    )
    parser.add_argument(
        '--sentry-dsn',
        help='Sentry DSN for error tracking (or set SENTRY_DSN env var)'
    )
    parser.add_argument(
        '--no-sentry',
        action='store_true',
        help='Disable Sentry error tracking'
    )
    parser.add_argument(
        '--db-path',
        default='results/file_organization.db',
        help='Path to SQLite database for persistent storage (default: results/file_organization.db)'
    )
    parser.add_argument(
        '--no-db',
        action='store_true',
        help='Disable database persistence (use in-memory registry only)'
    )
    parser.add_argument(
        '--run-migration',
        action='store_true',
        help='Run database migration to add canonical_id columns to existing records'
    )
    parser.add_argument(
        '--force',
        action='store_true',
        help='Force re-organization of all files, even if already in correct location'
    )

    args = parser.parse_args()

    # Initialize Sentry error tracking (before any other operations)
    if not args.no_sentry and ERROR_TRACKING_AVAILABLE:
        # Priority: CLI arg > FILE_SYSTEM_SENTRY_DSN > SENTRY_DSN
        sentry_dsn = args.sentry_dsn or os.environ.get('FILE_SYSTEM_SENTRY_DSN') or os.environ.get('SENTRY_DSN')
        if sentry_dsn:
            os.environ['SENTRY_DSN'] = sentry_dsn
        sentry_enabled = init_sentry()
        if sentry_enabled:
            print("✓ Sentry error tracking enabled")
    else:
        sentry_enabled = False

    # Run system health check
    if args.check_deps:
        from health_check import check_system
        check_system(verbose=True)
        return

    if not args.skip_health_check:
        from health_check import SystemHealthChecker
        checker = SystemHealthChecker().run_all_checks()
        checker.print_status()

    # Run migration if requested
    if args.run_migration:
        if GRAPH_STORE_AVAILABLE:
            from storage.migration import run_migration
            print(f"\n{'='*60}")
            print("Running ID Generation Migration")
            print(f"{'='*60}\n")
            run_migration(args.db_path)
            print("\nMigration complete. Canonical IDs have been generated for existing records.")
            return
        else:
            print("Error: GraphStore not available. Cannot run migration.")
            return

    # Create organizer with database path
    db_path = None if args.no_db else args.db_path
    organizer = ContentBasedFileOrganizer(
        base_path=args.base_path,
        enable_cost_tracking=not args.no_cost_tracking,
        db_path=db_path
    )

    # Organize directories
    summary = organizer.organize_directories(
        source_dirs=args.sources,
        dry_run=args.dry_run,
        limit=args.limit,
        force=args.force
    )

    # Print summary
    organizer.print_summary(summary)

    # Save reports
    if args.report or not args.dry_run:
        organizer.save_report(summary, args.report)

    # Save cost report if tracking was enabled
    if not args.no_cost_tracking and organizer.cost_calculator:
        organizer.save_cost_report(args.cost_report)

    # Update _site directory with latest HTML files
    if not args.dry_run:
        import subprocess
        from pathlib import Path
        script_path = Path(__file__).parent / 'copy_to_site.sh'
        if script_path.exists():
            try:
                subprocess.run([str(script_path)], check=True, capture_output=True)
                print("\n✓ Updated _site directory with latest HTML files")
            except subprocess.CalledProcessError:
                print("\n⚠ Failed to update _site directory")


if __name__ == '__main__':
    main()
