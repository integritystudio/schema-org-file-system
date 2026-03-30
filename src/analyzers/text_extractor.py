"""
Text extraction from images, PDFs, Word documents, and Excel spreadsheets.
"""

from __future__ import annotations

from contextlib import nullcontext
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional

# OCR (docTR via shared.ocr_utils) and PDF imports
try:
    from shared.ocr_utils import (
        extract_ocr_text,
        extract_ocr_with_confidence,
        extract_ocr_text_pdf,
        extract_ocr_pdf_with_confidence,
        is_ocr_available,
        OCRResult,
    )
    import pypdf
    OCR_AVAILABLE = is_ocr_available()
except ImportError:
    OCR_AVAILABLE = False

# Word document imports
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel imports
try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Cost tracking imports (optional)
try:
    from cost_roi_calculator import CostTracker
except ImportError:
    class CostTracker:  # type: ignore[no-redef]
        """Stub CostTracker when cost tracking is not available."""

        def __init__(self, *args: Any, **kwargs: Any) -> None:
            pass

        def __enter__(self) -> "CostTracker":
            return self

        def __exit__(self, *args: Any) -> bool:
            return False


_MAX_PDF_PAGES = 10
_MAX_PDF_OCR_PAGES = 5
_MIN_PDF_TEXT_LENGTH = 100
_MAX_XLSX_SHEETS = 5
_MAX_XLSX_ROWS = 100
_MAX_TEXT_BYTES = 50_000

_LOW_CONFIDENCE_THRESHOLD = 0.3


@dataclass
class ExtractionResult:
    """Text extraction result with optional OCR metadata."""
    text: str
    confidence: Optional[float] = None     # OCR confidence (0.0–1.0), None for non-OCR
    language: Optional[str] = None         # detected language code
    source: str = "unknown"                # "ocr", "pypdf", "docx", "xlsx", "text"


class TextExtractor:
    """Extract text content from various file formats."""

    def __init__(self, cost_calculator: Any | None = None) -> None:
        self.cost_calculator = cost_calculator
        self.ocr_available = OCR_AVAILABLE

    def extract_text_from_image(self, image_path: Path) -> str:
        """Extract text from image using docTR OCR."""
        if not self.ocr_available:
            return ""

        with CostTracker(self.cost_calculator, 'doctr_ocr') if self.cost_calculator else nullcontext():
            try:
                result = extract_ocr_text(image_path, max_chars=0)
                return result or ""
            except Exception as e:
                print(f"  OCR error: {e}")
                return ""

    def extract_from_image(self, image_path: Path) -> ExtractionResult:
        """Extract text from image with confidence metadata."""
        if not self.ocr_available:
            return ExtractionResult(text="", source="ocr")

        with CostTracker(self.cost_calculator, 'doctr_ocr') if self.cost_calculator else nullcontext():
            try:
                ocr_result = extract_ocr_with_confidence(image_path, max_chars=0)
                if ocr_result is None:
                    return ExtractionResult(text="", source="ocr")
                return ExtractionResult(
                    text=ocr_result.text,
                    confidence=ocr_result.confidence,
                    language=ocr_result.language,
                    source="ocr",
                )
            except Exception as e:
                print(f"  OCR error: {e}")
                return ExtractionResult(text="", source="ocr")

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF (searchable or scanned)."""
        if not self.ocr_available:
            return ""

        with CostTracker(self.cost_calculator, 'pdf_extraction') if self.cost_calculator else nullcontext():
            text = ""

            try:
                with open(pdf_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages[:_MAX_PDF_PAGES]:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"

                if len(text.strip()) > _MIN_PDF_TEXT_LENGTH:
                    return text.strip()

                print(f"  Using docTR OCR for scanned PDF...")
                ocr_text = extract_ocr_text_pdf(pdf_path, max_pages=_MAX_PDF_OCR_PAGES)
                if ocr_text:
                    text += ocr_text

                return text.strip()
            except Exception as e:
                print(f"  PDF extraction error: {e}")
                return ""

    def extract_from_pdf(self, pdf_path: Path) -> ExtractionResult:
        """Extract text from PDF with confidence metadata."""
        if not self.ocr_available:
            return ExtractionResult(text="", source="pypdf")

        with CostTracker(self.cost_calculator, 'pdf_extraction') if self.cost_calculator else nullcontext():
            text = ""

            try:
                with open(pdf_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages[:_MAX_PDF_PAGES]:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"

                if len(text.strip()) > _MIN_PDF_TEXT_LENGTH:
                    return ExtractionResult(
                        text=text.strip(),
                        confidence=1.0,    # searchable PDF text is exact
                        source="pypdf",
                    )

                print(f"  Using docTR OCR for scanned PDF...")
                ocr_result = extract_ocr_pdf_with_confidence(
                    pdf_path, max_pages=_MAX_PDF_OCR_PAGES
                )
                if ocr_result:
                    combined = (text + ocr_result.text).strip()
                    return ExtractionResult(
                        text=combined,
                        confidence=ocr_result.confidence,
                        language=ocr_result.language,
                        source="ocr",
                    )

                return ExtractionResult(text=text.strip(), source="pypdf")
            except Exception as e:
                print(f"  PDF extraction error: {e}")
                return ExtractionResult(text="", source="pypdf")

    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            return ""

        with CostTracker(self.cost_calculator, 'docx_extraction') if self.cost_calculator else nullcontext():
            try:
                doc = Document(docx_path)
                text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text.append(paragraph.text)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text.append(cell.text)

                return "\n".join(text)
            except Exception as e:
                print(f"  DOCX extraction error: {e}")
                return ""

    def extract_text_from_xlsx(self, xlsx_path: Path) -> str:
        """Extract text from Excel spreadsheet."""
        if not EXCEL_AVAILABLE:
            return ""

        with CostTracker(self.cost_calculator, 'xlsx_extraction') if self.cost_calculator else nullcontext():
            try:
                workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
                text = []

                for sheet_name in workbook.sheetnames[:_MAX_XLSX_SHEETS]:
                    sheet = workbook[sheet_name]
                    for row in sheet.iter_rows(max_row=_MAX_XLSX_ROWS, values_only=True):
                        row_text = ' '.join([str(cell) for cell in row if cell is not None])
                        if row_text.strip():
                            text.append(row_text)

                workbook.close()
                return "\n".join(text)
            except Exception as e:
                print(f"  XLSX extraction error: {e}")
                return ""

    def extract_text(self, file_path: Path, mime_type: str | None = None) -> str:
        """Extract text from various file types.

        Args:
            file_path: Path to the file.
            mime_type: Optional pre-detected MIME type. When omitted the
                       caller is responsible for passing it; pass an empty
                       string to skip MIME-based routing.
        """
        file_ext = file_path.suffix.lower()

        if mime_type and mime_type.startswith('image/'):
            return self.extract_text_from_image(file_path)
        elif mime_type == 'application/pdf' or file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self.extract_text_from_xlsx(file_path)
        elif (mime_type and mime_type.startswith('text/')) or file_ext in ['.txt', '.md', '.csv']:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(_MAX_TEXT_BYTES)
            except Exception:
                return ""

        return ""

    def extract(self, file_path: Path, mime_type: str | None = None) -> ExtractionResult:
        """Extract text with metadata from various file types.

        Enhanced version of extract_text() that returns confidence and language.
        """
        file_ext = file_path.suffix.lower()

        if mime_type and mime_type.startswith('image/'):
            return self.extract_from_image(file_path)
        elif mime_type == 'application/pdf' or file_ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return ExtractionResult(
                text=self.extract_text_from_docx(file_path),
                confidence=1.0,
                source="docx",
            )
        elif file_ext in ['.xlsx', '.xls']:
            return ExtractionResult(
                text=self.extract_text_from_xlsx(file_path),
                confidence=1.0,
                source="xlsx",
            )
        elif (mime_type and mime_type.startswith('text/')) or file_ext in ['.txt', '.md', '.csv']:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return ExtractionResult(
                        text=f.read(_MAX_TEXT_BYTES),
                        confidence=1.0,
                        source="text",
                    )
            except Exception:
                return ExtractionResult(text="", source="text")

        return ExtractionResult(text="", source="unknown")
